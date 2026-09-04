from __future__ import annotations

import pytest
from actweave_knowledge.extraction.builtin.markdown_extractor import markdown_sections
from actweave_knowledge.extraction.contracts import Document, SourceSpan
from actweave_knowledge.ingestion.index_text import build_index_text, has_indexable_source_text


def test_index_text_keeps_code_and_labels_but_not_attachment_uri() -> None:
    """Including Markdown destinations or dropping visible code breaks retrieval input."""

    attachment_ref = "a" * 64
    markdown = f"# 设备\n\n- 管理 IP：10.0.0.1\n\n```cpp\nList<int> values;\n```\n\n![端口照片](knowledge-attachment:{attachment_ref})"

    indexed = build_index_text(markdown)

    assert "设备" in indexed
    assert "管理 IP" in indexed and "10.0.0.1" in indexed
    assert "List<int> values;" in indexed
    assert "端口照片" in indexed
    assert "knowledge-attachment:" not in indexed and attachment_ref not in indexed


def test_index_text_keeps_table_link_label_and_literal_html_like_text() -> None:
    """Dropping visible table/link/literal tokens or regex-stripping angle text breaks search."""

    indexed = build_index_text("| 字段 | 值 |\n| --- | --- |\n| IP | <IP> |\n\n[设备说明](https://example.invalid/path?tracking=1)\n\n<script>alert('literal text')</script>")

    assert "字段" in indexed and "值" in indexed and "IP" in indexed
    assert "设备说明" in indexed
    assert "https://example.invalid" not in indexed
    assert "<IP>" in indexed
    assert "alert('literal text')" in indexed


def test_source_spans_ignore_context_prefix_generated_placeholder_and_image_alt() -> None:
    """Treating derived text or image alt as source makes scans falsely searchable."""

    image_markdown = "![本页图片](knowledge-attachment:" + "a" * 64 + ")"
    image_only = Document(
        page_content=image_markdown,
        source_spans=(SourceSpan(block_id="page:1:image:1", start=0, end=len(image_markdown), location={"page": 1}),),
    )
    context_only = Document(
        page_content="# 上下文标题\n\n[文本提取失败]",
        source_spans=(SourceSpan(block_id="heading", start=0, end=8, location={"paragraph": 1}, role="context_prefix"),),
    )
    external_image_only = tuple(markdown_sections("![端口图片](https://example.invalid/port.png)\n"))
    real_source = Document(
        page_content="# 上下文标题\n\n真实设备编号 R-01",
        source_spans=(SourceSpan(block_id="body", start=10, end=20, location={"paragraph": 2}),),
    )

    assert not has_indexable_source_text((image_only,))
    assert not has_indexable_source_text((context_only,))
    assert not has_indexable_source_text(external_image_only)
    assert has_indexable_source_text((real_source,))


def test_physical_line_source_spans_keep_fenced_markdown_literals_indexable() -> None:
    """Reparsing a physical source line loses its enclosing fence semantics."""

    documents = tuple(markdown_sections("```markdown\n![diagram](https://example.invalid/diagram.png)\n[router]: https://example.invalid/router\n```\n"))

    assert build_index_text(documents[0].page_content).strip()
    assert has_indexable_source_text(documents)


def test_physical_line_source_spans_keep_indented_multiline_code_indexable() -> None:
    """Searching de-indented token content in raw Markdown loses real code coverage."""

    documents = tuple(markdown_sections("    print('one')\n    print('two')\n"))

    assert "print('one')" in build_index_text(documents[0].page_content)
    assert has_indexable_source_text(documents)


def test_inline_wrappers_cannot_make_a_generated_image_placeholder_indexable() -> None:
    """Visible placeholder text must not inherit authority from source formatting syntax."""

    image = "![diagram](https://example.invalid/diagram.png)"
    documents = tuple(tuple(markdown_sections(markdown)) for markdown in (f"_{image}_\n", f"**{image}**\n", f"[{image}](https://example.invalid/docs)\n"))

    assert all("外部图片未获取：diagram" in item[0].page_content for item in documents)
    assert all(not has_indexable_source_text(item) for item in documents)


@pytest.mark.parametrize(
    ("markdown", "visible"),
    [
        (r"设备 \<IP\>", "设备 <IP>"),
        ("router &amp; switch", "router & switch"),
        ("`router\nconfigure`", "router configure"),
        ("&#33;", "!"),
    ],
)
def test_visible_source_survives_markdown_decoding_and_code_normalization(markdown: str, visible: str) -> None:
    """Locating decoded tokens by substring rejects real source text and punctuation."""

    documents = tuple(markdown_sections(markdown))

    assert build_index_text(documents[0].page_content) == visible
    assert has_indexable_source_text(documents)


@pytest.mark.parametrize("separator", ["\n", " &nbsp; "])
def test_source_whitespace_between_generated_images_cannot_supply_text(separator: str) -> None:
    """Source whitespace must not confer text authority on neighboring generated nodes."""

    documents = tuple(markdown_sections(f"![a](https://example.invalid/a.png){separator}![b](https://example.invalid/b.png)\n"))

    assert "外部图片未获取：a" in documents[0].page_content
    assert "外部图片未获取：b" in documents[0].page_content
    assert not has_indexable_source_text(documents)


def test_word_table_escaped_punctuation_keeps_its_cell_source(tmp_path) -> None:
    """Table de-escaping must not attribute visible cell text to generated separators."""

    from actweave_knowledge.extraction.processor import ExtractProcessor
    from docx import Document as WordFile
    from docx.oxml import OxmlElement
    from parsing_test_helpers import make_context, make_setting

    path = tmp_path / "source.docx"
    word = WordFile()
    table = word.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "|"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    word.save(path)
    documents = tuple(ExtractProcessor().extract(make_setting(path), make_context(tmp_path / "work")))

    assert build_index_text(documents[0].page_content) == "|"
    assert has_indexable_source_text(documents)


def test_real_punctuation_beside_a_generated_image_remains_indexable() -> None:
    """A mixed generated/source inline must keep the source's own visible characters."""

    documents = tuple(markdown_sections("![a](https://example.invalid/a.png) &#33;"))

    assert has_indexable_source_text(documents)


@pytest.mark.parametrize(
    "markdown",
    [
        "[![a](https://example.invalid/a.png)][docs]\n\n[docs]: https://example.invalid/docs",
        "| _![a](https://example.invalid/a.png)_ | **![b](https://example.invalid/b.png)** |\n| --- | --- |",
    ],
)
def test_reference_and_table_syntax_cannot_authorize_generated_images(markdown: str) -> None:
    """Reference destinations and table delimiters are source syntax, not visible text."""

    documents = tuple(markdown_sections(markdown))

    assert "外部图片未获取" in documents[0].page_content
    assert not has_indexable_source_text(documents)


def test_html_generated_image_placeholders_do_not_supply_source_text(tmp_path) -> None:
    """HTML-generated placeholders stay excluded even with source emphasis and spacing."""

    from actweave_knowledge.extraction.processor import ExtractProcessor
    from parsing_test_helpers import make_context, make_setting

    path = tmp_path / "source.html"
    path.write_text('<p><em><img src="https://example.invalid/a.png" alt="a"></em> <img src="https://example.invalid/b.png" alt="b"></p>', encoding="utf-8")
    documents = tuple(ExtractProcessor().extract(make_setting(path), make_context(tmp_path / "work")))

    assert "外部图片未获取" in documents[0].page_content
    assert not has_indexable_source_text(documents)

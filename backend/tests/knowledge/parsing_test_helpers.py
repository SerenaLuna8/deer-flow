from __future__ import annotations

import hashlib
import io
from pathlib import Path

from actweave_knowledge.extraction.contracts import (
    Attachment,
    ChunkProfile,
    Document,
    ExtractionContext,
    ExtractionLimits,
    ExtractSetting,
    LocalAttachment,
    ParseProfile,
    SourceSpan,
)


def make_parse_profile(extension: str, *, etl_type: str = "dify", header_rules: tuple = ()) -> ParseProfile:
    from actweave_knowledge.extraction.registry import default_registry

    item = default_registry().resolve(
        datasource_type="file",
        etl_type=etl_type,
        extension=extension,
    )
    return ParseProfile(
        etl_type=etl_type,
        extractor_id=item.extractor_id,
        extractor_version=item.extractor_version,
        normalization_version="md-v1",
        image_policy_version="raster-v1",
        header_rules=tuple(header_rules),
    )


def make_chunk_profile(**overrides: object) -> ChunkProfile:
    fields: dict[str, object] = {
        "unit": "token",
        "mode": "general",
        "size": 1000,
        "overlap": 100,
        "separator": "\n\n",
        "child_size": 500,
        "child_separator": "\n\n",
        "remove_extra_spaces": False,
        "remove_urls_emails": False,
        "tokenizer_profile_id": "knowledge-cl100k-v1",
        "tokenizer_digest": None,
        "cleaner_version": "cleaner-v1",
        "splitter_version": "splitter-v1",
    }
    fields.update(overrides)
    if fields["unit"] == "token":
        from actweave_knowledge.ingestion.tokenizer import TOKENIZER_PROFILE_ID, tokenizer_fingerprint

        fields["tokenizer_profile_id"] = TOKENIZER_PROFILE_ID
        fields["tokenizer_digest"] = tokenizer_fingerprint()
    else:
        fields["tokenizer_profile_id"] = None
        fields["tokenizer_digest"] = None
    return ChunkProfile(**fields)


def make_document(text: str, *, location: dict[str, str | int] | None = None, heading_path: tuple[str, ...] = ()) -> Document:
    return Document(
        page_content=text,
        source_spans=(
            SourceSpan(
                block_id="block:1",
                start=0,
                end=len(text),
                location=location or {"paragraph": 1},
            ),
        ),
        heading_path=tuple(heading_path),
        kind="paragraph",
    )


def make_setting(path: str | Path, **overrides: object) -> ExtractSetting:
    source_path = Path(path)
    fields: dict[str, object] = {
        "source_path": source_path,
        "original_name": source_path.name,
        "datasource_type": "file",
        "profile": overrides["profile"] if "profile" in overrides else make_parse_profile(source_path.suffix),
    }
    fields.update(overrides)
    return ExtractSetting(**fields)


class CollectingAttachmentSink:
    def __init__(self, work_dir: Path) -> None:
        self.work_dir = work_dir
        self.assets: list[LocalAttachment] = []
        self.occurrences: list[tuple[str, str, SourceSpan]] = []

    def accept(self, source_path: Path, *, alt_text: str, source: SourceSpan) -> Attachment:
        from PIL import Image

        data = source_path.read_bytes()
        ref = hashlib.sha256(data).hexdigest()
        with Image.open(source_path) as image:
            width, height = image.size
        attachment = Attachment(
            ref=ref,
            media_type="image/png",
            size_bytes=len(data),
            width=width,
            height=height,
        )
        target = self.work_dir / f"{ref}.png"
        target.write_bytes(data)
        if not any(item.attachment.ref == ref for item in self.assets):
            self.assets.append(LocalAttachment(attachment=attachment, relative_path=target.name))
        self.occurrences.append((ref, alt_text, source))
        return attachment


def make_context(work_dir: Path) -> ExtractionContext:
    work_dir.mkdir(parents=True, exist_ok=True)
    return ExtractionContext(
        work_dir=work_dir,
        sink=CollectingAttachmentSink(work_dir),
        limits=ExtractionLimits(),
        check_cancelled=lambda: None,
    )


def write_pdf(path: Path, pages: list[str]) -> None:
    """Public fixture migrated from test_ingestion._write_pdf (Schema V1 gates)."""

    page_entries: list[tuple[int, int, str]] = []
    next_id = 3
    for line in pages:
        page_entries.append((next_id, next_id + 1, line))
        next_id += 2
    font_id = next_id
    kids = " ".join(f"{page_id} 0 R" for page_id, _, _ in page_entries)
    objects: list[tuple[int, bytes]] = [
        (1, b"<< /Type /Catalog /Pages 2 0 R >>"),
        (2, f"<< /Type /Pages /Kids [{kids}] /Count {len(page_entries)} >>".encode()),
    ]
    for page_id, content_id, line in page_entries:
        objects.append(
            (
                page_id,
                (f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents {content_id} 0 R /Resources << /Font << /F1 {font_id} 0 R >> >> >>").encode(),
            )
        )
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
        objects.append(
            (
                content_id,
                b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
            )
        )
    objects.append((font_id, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))

    output = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for object_id, body in objects:
        offsets[object_id] = len(output)
        output += f"{object_id} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_offset = len(output)
    output += f"xref\n0 {len(objects) + 1}\n".encode()
    output += b"0000000000 65535 f \n"
    for object_id in sorted(offsets):
        output += f"{offsets[object_id]:010d} 00000 n \n".encode()
    output += (f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n").encode()
    path.write_bytes(bytes(output))


def write_docx_with_image(path: Path) -> None:
    """Write one deterministic Word document with an embedded local raster."""

    from docx import Document as WordFile
    from PIL import Image

    document = WordFile()
    document.add_heading("设备手册", level=1)
    document.add_paragraph("管理接口地址为 10.0.0.1，请检查链路状态。")
    with io.BytesIO() as image_bytes:
        with Image.new("RGB", (8, 8), "red") as image:
            image.save(image_bytes, format="PNG")
        image_bytes.seek(0)
        document.add_picture(image_bytes)
        document.save(path)

"""P3-T5 gates for the P1/P2/P3 production ingestion boundary."""

from __future__ import annotations

import asyncio
import uuid
from pathlib import Path

import pytest
from actweave_knowledge import KNOWLEDGE_QUOTA_EXCEEDED, KnowledgeError
from actweave_knowledge.extraction.contracts import ProcessingProfile
from actweave_knowledge.ingestion import pipeline as pipeline_module
from actweave_knowledge.persistence.models import (
    KnowledgeAttachmentRow,
    KnowledgeBaseRow,
    KnowledgeDocumentRow,
    KnowledgeExtractionRow,
    KnowledgeSegmentRow,
)
from extraction_test_helpers import extraction_harness
from ingestion_test_helpers import ingestion_harness
from parsing_test_helpers import make_chunk_profile, make_parse_profile, write_docx_with_image
from sqlalchemy import event, select, text
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import Session


@pytest.mark.asyncio
async def test_abort_incomplete_extraction_atomically_unpins_and_enqueues_cleanup(
    postgres_database_url: str,
) -> None:
    """A parser failure transfers one exact staging generation to durable GC."""

    async with extraction_harness(postgres_database_url) as resources:
        rows = await resources.read_rows()
        source_sha256 = rows["documents"][0].source_sha256
        assert source_sha256 is not None
        reservation = await resources.store.begin(
            resources.claim,
            source_sha256=source_sha256,
            profile=make_parse_profile(".pdf"),
        )

        await resources.store.abort(resources.claim, reservation)
        # A callback may have transferred the same generation first.  The
        # parser owner's final abort is an exact-claim idempotent no-op.
        await resources.store.abort(resources.claim, reservation)

        rows = await resources.read_rows()
        extraction = next(row for row in rows["extractions"] if row.id == reservation.extraction_id)
        source_task = next(row for row in rows["tasks"] if row.id == resources.claim.id)
        cleanup = [row for row in rows["tasks"] if row.kind == "delete_extraction"]
        assert extraction.state == "deleting"
        assert source_task.extraction_id is None
        assert len(cleanup) == 1
        assert cleanup[0].resource_id == reservation.extraction_id


@pytest.mark.asyncio
async def test_ingest_matches_preview_and_embeds_index_text(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """Formal ingestion consumes the same frozen P1/P3 result as preview."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "manual.docx"
        write_docx_with_image(source)
        profile = ProcessingProfile(
            parse=make_parse_profile(".docx"),
            chunk=make_chunk_profile(),
        )

        preview = await harness.preview(source, profile)
        uploaded = await harness.upload(source, profile)
        await harness.run_next_task()
        rows = await harness.segments(uploaded.id)

        assert [row.content for row in rows[:10]] == [chunk.content for chunk in preview.chunks]
        assert harness.fake_model.calls[-1] == [row.index_text for row in rows]
        assert all("knowledge-attachment:" not in value for value in harness.fake_model.calls[-1])
        assert all(row.token_count > 0 and row.source_spans for row in rows)
        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        assert document.published_extraction_id == rows[0].extraction_id
        assert document.parsing_profile == profile.model_dump(mode="json")
        assert len(facts["attachments"]) == 1
        assert len(facts["bindings"]) == 1
        assert facts["bindings"][0].extraction_id == document.published_extraction_id


@pytest.mark.asyncio
async def test_embedding_failure_preserves_ready_cache_and_retry_hits_it(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """Provider failure keeps a complete pinned-then-released Extraction."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "cache.md"
        source.write_text("# 缓存\n\n正文内容与 [链接](https://example.invalid/x)。", encoding="utf-8")
        profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        harness.fake_model.fail = True

        failed = await harness.run_next_task(expected_status="retry_wait")

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        extraction = facts["extractions"][0]
        assert extraction.state == "ready"
        assert document.published_extraction_id is None
        assert failed.extraction_id is None
        assert not [row for row in facts["tasks"] if row.kind == "delete_extraction"]
        source_gets = sum(operation == "get" and key == document.storage_key for operation, key in harness.resources.object_store.calls)

        harness.fake_model.fail = False
        await harness.run_next_task()

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        assert len(facts["extractions"]) == 1
        assert document.published_extraction_id == extraction.id
        assert source_gets == sum(operation == "get" and key == document.storage_key for operation, key in harness.resources.object_store.calls)


@pytest.mark.asyncio
async def test_parser_failure_atomically_enqueues_only_its_incomplete_generation(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """Invalid source bytes never leave a pinned or visible staging result."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "broken.pdf"
        source.write_bytes(b"not a pdf")
        profile = ProcessingProfile(
            parse=make_parse_profile(".pdf"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)

        failed = await harness.run_next_task(expected_status="retry_wait")

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        extraction = facts["extractions"][0]
        cleanup = [row for row in facts["tasks"] if row.kind == "delete_extraction"]
        assert document.published_extraction_id is None
        assert extraction.state == "deleting"
        assert failed.extraction_id is None
        assert len(cleanup) == 1 and cleanup[0].resource_id == extraction.id


@pytest.mark.asyncio
async def test_attachment_put_failure_keeps_publication_empty_and_cleanup_durable(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """An on_asset infrastructure failure cannot degrade to text-only."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "asset.docx"
        write_docx_with_image(source)
        profile = ProcessingProfile(
            parse=make_parse_profile(".docx"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        harness.resources.object_store.fail_next("put")

        failed = await harness.run_next_task(expected_status="retry_wait")

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        cleanup = [row for row in facts["tasks"] if row.kind == "delete_extraction"]
        assert document.published_extraction_id is None
        assert await harness.segments(uploaded.id) == []
        assert failed.extraction_id is None
        assert len(cleanup) == 1
        assert facts["extractions"][0].state == "deleting"


@pytest.mark.asyncio
async def test_manifest_put_starts_after_parser_workspace_is_removed(
    postgres_database_url: str,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Store.complete owns the only live temp directory at manifest PUT."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "manifest.md"
        source.write_text("# Manifest\n\n正文", encoding="utf-8")
        profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        workspaces: list[Path] = []
        real_mkdtemp = pipeline_module.tempfile.mkdtemp

        def tracked_mkdtemp(*args, **kwargs):  # noqa: ANN002, ANN003
            path = Path(real_mkdtemp(*args, **kwargs))
            if kwargs.get("prefix") == "actweave-knowledge-ingest-":
                workspaces.append(path)
            return str(path)

        real_upload = harness.resources.object_store.upload_from

        async def checked_upload(key: str, path: Path, *, media_type: str | None = None) -> None:
            if key.endswith("/manifest.json"):
                assert workspaces and not any(item.exists() for item in workspaces)
            await real_upload(key, path, media_type=media_type)

        monkeypatch.setattr(pipeline_module.tempfile, "mkdtemp", tracked_mkdtemp)
        monkeypatch.setattr(harness.resources.object_store, "upload_from", checked_upload)
        harness.resources.object_store.fail_next("put")

        failed = await harness.run_next_task(expected_status="retry_wait")

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        assert document.published_extraction_id is None
        assert facts["extractions"][0].state == "deleting"
        assert failed.extraction_id is None
        assert workspaces and not any(path.exists() for path in workspaces)


@pytest.mark.asyncio
async def test_manifest_encoding_failure_aborts_the_incomplete_generation(
    postgres_database_url: str,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A failure before manifest registration still enters durable cleanup."""

    from actweave_knowledge.storage import extractions as extractions_module

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "encode.md"
        source.write_text("# Encode\n\n正文", encoding="utf-8")
        profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)

        def fail_encode(_result):  # noqa: ANN001, ANN202
            raise KnowledgeError(KNOWLEDGE_QUOTA_EXCEEDED, "manifest fixture")

        monkeypatch.setattr(extractions_module, "encode_manifest", fail_encode)

        failed = await harness.run_next_task(expected_status="retry_wait")

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        cleanup = [row for row in facts["tasks"] if row.kind == "delete_extraction"]
        assert document.published_extraction_id is None
        assert facts["extractions"][0].state == "deleting"
        assert len(cleanup) == 1
        assert failed.extraction_id is None


@pytest.mark.asyncio
async def test_publish_flush_failure_rolls_back_rows_pointer_and_keeps_ready_cache(
    postgres_database_url: str,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A publication flush error exposes neither half of the new generation."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "flush.md"
        source.write_text("# Flush\n\n原子发布正文", encoding="utf-8")
        profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        real_flush = AsyncSession.flush
        injected = False

        async def fail_segment_flush(self, objects=None):  # noqa: ANN001
            nonlocal injected
            if not injected and any(isinstance(row, KnowledgeSegmentRow) for row in self.new):
                injected = True
                raise SQLAlchemyError("injected publish flush failure")
            return await real_flush(self, objects)

        monkeypatch.setattr(AsyncSession, "flush", fail_segment_flush)

        failed = await harness.run_next_task(expected_status="retry_wait")

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        assert injected is True
        assert document.published_extraction_id is None
        assert await harness.segments(uploaded.id) == []
        assert facts["extractions"][0].state == "ready"
        assert failed.extraction_id is None


@pytest.mark.asyncio
async def test_document_version_race_after_embedding_publishes_nothing(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """A CAS loser settles without replacing content or the Extraction pointer."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "race.md"
        source.write_text("# Race\n\n版本竞争正文", encoding="utf-8")
        profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        blocker = asyncio.Event()
        harness.fake_model.blocker = blocker
        running = asyncio.create_task(harness.run_next_task())
        await asyncio.wait_for(harness.fake_model.started.wait(), timeout=10)
        async with harness.resources.session_factory() as session, session.begin():
            document = await session.get(KnowledgeDocumentRow, uploaded.id, with_for_update=True)
            assert document is not None
            document.version += 1
            document.status = "deleting"
        blocker.set()

        await running

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        assert document.status == "deleting"
        assert document.published_extraction_id is None
        assert await harness.segments(uploaded.id) == []


@pytest.mark.asyncio
async def test_changed_source_bytes_abort_before_parser_and_preserve_old_publication(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """Downloaded bytes must match the admission-frozen source digest."""

    async with ingestion_harness(
        postgres_database_url,
        cache_enabled=False,
    ) as harness:
        source = tmp_path / "source.md"
        source.write_text("# Original\n\n可信正文", encoding="utf-8")
        profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        await harness.run_next_task()
        original_segments = await harness.segments(uploaded.id)
        facts = await harness.resources.read_rows()
        original_document = next(row for row in facts["documents"] if row.id == uploaded.id)
        original_extraction = original_document.published_extraction_id
        assert original_extraction is not None

        harness.resources.object_store.objects[original_document.storage_key] = b"tampered bytes"
        harness.fake_model.calls.clear()
        await harness.reparse(uploaded.id, profile)

        failed = await harness.run_next_task(expected_status="retry_wait")

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        newer = [row for row in facts["extractions"] if row.id != original_extraction]
        cleanup = [row for row in facts["tasks"] if row.kind == "delete_extraction"]
        assert document.published_extraction_id == original_extraction
        assert [row.id for row in await harness.segments(uploaded.id)] == [row.id for row in original_segments]
        assert len(newer) == 1 and newer[0].state == "deleting"
        assert len(cleanup) == 1 and cleanup[0].resource_id == newer[0].id
        assert failed.extraction_id is None
        assert harness.fake_model.calls == []


@pytest.mark.asyncio
async def test_embedding_barrier_lease_loss_never_reaches_publication(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """The post-Provider guard stops a response owned by an expired claim."""

    from actweave_knowledge.persistence.tasks import recover_expired_tasks

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "lease.md"
        source.write_text("# Lease\n\n租约边界正文", encoding="utf-8")
        profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        blocker = asyncio.Event()
        harness.fake_model.blocker = blocker
        running = asyncio.create_task(
            harness.run_next_task(expected_status=None),
        )
        await asyncio.wait_for(harness.fake_model.started.wait(), timeout=10)
        facts = await harness.resources.read_rows()
        ingest_task = next(row for row in facts["tasks"] if row.kind == "ingest_document" and row.resource_id == uploaded.id)
        async with harness.resources.session_factory() as session, session.begin():
            await session.execute(
                text("UPDATE knowledge_tasks SET lease_until = clock_timestamp() - interval '1 second' WHERE id = :task_id"),
                {"task_id": ingest_task.id},
            )
        blocker.set()

        await running

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        source_task = next(row for row in facts["tasks"] if row.id == ingest_task.id)
        assert source_task.status == "running"
        assert document.published_extraction_id is None
        assert await harness.segments(uploaded.id) == []
        async with harness.resources.session_factory() as session, session.begin():
            assert await recover_expired_tasks(session) >= 1
        facts = await harness.resources.read_rows()
        source_task = next(row for row in facts["tasks"] if row.id == ingest_task.id)
        assert source_task.status == "retry_wait" and source_task.extraction_id is None
        assert facts["extractions"][0].state == "ready"


@pytest.mark.asyncio
async def test_final_commit_failure_preserves_complete_old_generation(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """A commit failure rolls back rows, bindings, pointer and Task settlement."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "commit.docx"
        write_docx_with_image(source)
        first_profile = ProcessingProfile(
            parse=make_parse_profile(".docx"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, first_profile)
        await harness.run_next_task()
        old_segments = await harness.segments(uploaded.id)
        old_facts = await harness.resources.read_rows()
        old_document = next(row for row in old_facts["documents"] if row.id == uploaded.id)
        old_extraction = old_document.published_extraction_id
        old_binding_ids = [(row.segment_id, row.position, row.attachment_id) for row in old_facts["bindings"]]
        second_profile = ProcessingProfile(
            parse=first_profile.parse,
            chunk=make_chunk_profile(size=200, overlap=0),
        )
        await harness.reparse(uploaded.id, second_profile)
        injected = False

        def fail_publish_commit(session: Session) -> None:
            nonlocal injected
            if not injected and any(isinstance(row, KnowledgeDocumentRow) and row.version == 2 and row.published_version == 2 and row.status == "ready" for row in session.identity_map.values()):
                injected = True
                raise SQLAlchemyError("injected final commit failure")

        event.listen(Session, "before_commit", fail_publish_commit)
        try:
            failed = await harness.run_next_task(expected_status="retry_wait")
        finally:
            event.remove(Session, "before_commit", fail_publish_commit)

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        assert injected is True
        assert document.published_extraction_id == old_extraction
        assert document.published_version == 1
        assert [row.id for row in await harness.segments(uploaded.id)] == [row.id for row in old_segments]
        assert [(row.segment_id, row.position, row.attachment_id) for row in facts["bindings"]] == old_binding_ids
        assert failed.extraction_id is None


@pytest.mark.asyncio
async def test_missing_ready_manifest_reextracts_without_half_publication(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """Confirmed missing cache bytes are a miss under the live claim."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "missing.md"
        source.write_text("# Cache\n\n缓存损坏后重新解析。", encoding="utf-8")
        first_profile = ProcessingProfile(
            parse=make_parse_profile(".md"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, first_profile)
        await harness.run_next_task()
        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        old_extraction = next(row for row in facts["extractions"] if row.id == document.published_extraction_id)
        assert old_extraction.manifest_storage_key is not None
        harness.resources.object_store.objects.pop(old_extraction.manifest_storage_key)
        source_gets = sum(operation == "get" and key == document.storage_key for operation, key in harness.resources.object_store.calls)
        second_profile = ProcessingProfile(
            parse=first_profile.parse,
            chunk=make_chunk_profile(size=200, overlap=0),
        )
        await harness.reparse(uploaded.id, second_profile)

        await harness.run_next_task()

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        assert document.published_extraction_id != old_extraction.id
        assert len(facts["extractions"]) == 2
        assert source_gets + 1 == sum(operation == "get" and key == document.storage_key for operation, key in harness.resources.object_store.calls)
        rows = await harness.segments(uploaded.id)
        assert rows and all(row.extraction_id == document.published_extraction_id for row in rows)


@pytest.mark.asyncio
async def test_repeated_image_occurrences_share_bytes_but_keep_ordered_bindings(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """One normalized Attachment may appear more than once in one Segment."""

    import docx
    from PIL import Image

    image_path = tmp_path / "repeat.png"
    Image.new("RGB", (4, 4), "blue").save(image_path, format="PNG")
    document = docx.Document()
    document.add_heading("重复图片", level=1)
    document.add_paragraph("同一图片在正文中出现两次。")
    document.add_picture(str(image_path))
    document.add_picture(str(image_path))
    source = tmp_path / "repeat.docx"
    document.save(source)
    profile = ProcessingProfile(
        parse=make_parse_profile(".docx"),
        chunk=make_chunk_profile(),
    )

    async with ingestion_harness(postgres_database_url) as harness:
        uploaded = await harness.upload(source, profile)
        await harness.run_next_task()

        facts = await harness.resources.read_rows()
        document_row = next(row for row in facts["documents"] if row.id == uploaded.id)
        bindings = sorted(facts["bindings"], key=lambda row: row.position)
        assert len(facts["attachments"]) == 1
        assert len(bindings) == 2
        assert [row.position for row in bindings] == [1, 2]
        assert bindings[0].attachment_id == bindings[1].attachment_id
        assert all(row.extraction_id == document_row.published_extraction_id for row in bindings)


@pytest.mark.asyncio
async def test_attachment_manifest_mismatch_rolls_back_publication(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """Publication revalidates the locked Attachment inventory against manifest."""

    async with ingestion_harness(postgres_database_url) as harness:
        source = tmp_path / "mismatch.docx"
        write_docx_with_image(source)
        profile = ProcessingProfile(
            parse=make_parse_profile(".docx"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        blocker = asyncio.Event()
        harness.fake_model.blocker = blocker
        running = asyncio.create_task(
            harness.run_next_task(expected_status=None),
        )
        await asyncio.wait_for(harness.fake_model.started.wait(), timeout=10)
        async with harness.resources.session_factory() as session, session.begin():
            attachment = await session.scalar(select(KnowledgeAttachmentRow).with_for_update())
            assert attachment is not None
            attachment.width += 1
        blocker.set()

        await running

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        task = next(row for row in facts["tasks"] if row.kind == "ingest_document" and row.resource_id == uploaded.id)
        assert document.published_extraction_id is None
        assert task.status == "retry_wait" and task.extraction_id is None
        assert await harness.segments(uploaded.id) == []


@pytest.mark.asyncio
async def test_summary_model_lock_wait_past_deadline_rolls_back_publication(
    postgres_database_url: str,
    tmp_path: Path,
) -> None:
    """The final DB-clock fence follows every wait-capable summary operation."""

    import docx
    from actweave_knowledge.persistence.tasks import recover_expired_tasks
    from support.system_model_seed import seed_system_model_config

    from app.knowledge_settings.service import default_knowledge_settings_row
    from deerflow.persistence.knowledge_settings import KnowledgeSystemSettingsRow
    from deerflow.persistence.system_settings import SystemModelConfigRow

    async with ingestion_harness(
        postgres_database_url,
        cache_enabled=False,
    ) as harness:
        source = tmp_path / "summary-lock.docx"
        write_docx_with_image(source)
        document_file = docx.Document(source)
        document_file.add_paragraph("长期正文" * 100)
        document_file.save(source)
        profile = ProcessingProfile(
            parse=make_parse_profile(".docx"),
            chunk=make_chunk_profile(),
        )
        uploaded = await harness.upload(source, profile)
        await harness.run_next_task()
        old_segments = await harness.segments(uploaded.id)
        old_facts = await harness.resources.read_rows()
        old_document = next(row for row in old_facts["documents"] if row.id == uploaded.id)
        old_extraction_id = old_document.published_extraction_id
        old_bindings = [(row.segment_id, row.position, row.attachment_id) for row in old_facts["bindings"]]
        assert old_extraction_id is not None and old_bindings

        summary_model_id = uuid.uuid4()
        async with harness.resources.session_factory() as session, session.begin():
            owner_user_id = await session.scalar(
                text("SELECT created_by_user_id FROM projects WHERE id = :project_id"),
                {"project_id": harness.resources.project_id},
            )
            assert owner_user_id is not None
            await seed_system_model_config(
                session,
                model_id=summary_model_id,
                owner_user_id=str(owner_user_id),
                display_name="Summary lock model",
                provider_model="summary/lock-test",
            )
            settings = await session.get(KnowledgeSystemSettingsRow, 1)
            if settings is None:
                settings = default_knowledge_settings_row()
                session.add(settings)
            settings.summary_model_name = str(summary_model_id)
            base = await session.get(
                KnowledgeBaseRow,
                harness.resources.base_id,
                with_for_update=True,
            )
            assert base is not None
            base.summary_index_enabled = True

        await harness.reparse(uploaded.id, profile)
        model_locker = harness.resources.session_factory()
        await model_locker.begin()
        await model_locker.execute(select(SystemModelConfigRow).where(SystemModelConfigRow.id == summary_model_id).with_for_update())
        embedding_release = asyncio.Event()
        harness.fake_model.blocker = embedding_release
        harness.fake_model.started.clear()
        running = asyncio.create_task(
            harness.run_next_task(expected_status=None),
        )
        try:
            await asyncio.wait_for(harness.fake_model.started.wait(), timeout=10)
            facts = await harness.resources.read_rows()
            ingest_task = next(row for row in facts["tasks"] if row.kind == "ingest_document" and row.resource_id == uploaded.id and row.status == "running")
            async with harness.resources.session_factory() as session, session.begin():
                await session.execute(
                    text("UPDATE knowledge_tasks SET lease_until = clock_timestamp() + interval '1 second' WHERE id = :task_id"),
                    {"task_id": ingest_task.id},
                )
            embedding_release.set()

            async with asyncio.timeout(10):
                while True:
                    async with harness.resources.session_factory() as session:
                        waiting = await session.scalar(
                            text("SELECT EXISTS (SELECT 1 FROM pg_stat_activity WHERE datname = current_database() AND pid <> pg_backend_pid() AND wait_event_type = 'Lock' AND query ILIKE '%system_model_configs%')")
                        )
                    if waiting:
                        break
                    await asyncio.sleep(0.01)

            async with asyncio.timeout(10):
                while True:
                    async with harness.resources.session_factory() as session:
                        expired = await session.scalar(
                            text("SELECT clock_timestamp() > lease_until FROM knowledge_tasks WHERE id = :task_id"),
                            {"task_id": ingest_task.id},
                        )
                    if expired:
                        break
                    await asyncio.sleep(0.01)
        finally:
            embedding_release.set()
            await model_locker.rollback()
            await model_locker.close()

        await running

        facts = await harness.resources.read_rows()
        document = next(row for row in facts["documents"] if row.id == uploaded.id)
        source_task = next(row for row in facts["tasks"] if row.id == ingest_task.id)
        assert document.published_extraction_id == old_extraction_id
        assert document.published_version == 1
        assert [row.id for row in await harness.segments(uploaded.id)] == [row.id for row in old_segments]
        assert [(row.segment_id, row.position, row.attachment_id) for row in facts["bindings"]] == old_bindings
        assert source_task.status == "running" and source_task.extraction_id is not None
        assert not [row for row in facts["tasks"] if row.kind == "summarize_document" and row.status != "succeeded"]
        new_extractions = [row for row in facts["extractions"] if isinstance(row, KnowledgeExtractionRow) and row.id != old_extraction_id]
        assert len(new_extractions) == 1 and new_extractions[0].state == "ready"

        async with harness.resources.session_factory() as session, session.begin():
            assert await recover_expired_tasks(session) >= 1
        facts = await harness.resources.read_rows()
        source_task = next(row for row in facts["tasks"] if row.id == ingest_task.id)
        assert source_task.status == "retry_wait" and source_task.extraction_id is None

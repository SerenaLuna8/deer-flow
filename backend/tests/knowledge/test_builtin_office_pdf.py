"""Real local Office/PDF fixtures for the fixed-source upstream adapters."""

from __future__ import annotations

import hashlib

import pytest
from actweave_knowledge.contracts import KNOWLEDGE_QUOTA_EXCEEDED, KnowledgeError
from actweave_knowledge.extraction.contracts import ExtractionLimits, ExtractionResult
from actweave_knowledge.extraction.images import LocalAttachmentSink
from actweave_knowledge.extraction.manifest import canonical_parse_fingerprint, decode_manifest, encode_manifest
from actweave_knowledge.extraction.processor import ExtractProcessor
from docx import Document as WordFile
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from parsing_test_helpers import make_context, make_setting, write_pdf
from PIL import Image
from pypdf import PdfReader, PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject, NumberObject


def _extract(path, work_dir, *, limits=None, check_cancelled=None):
    context = make_context(work_dir)
    limits = limits or context.limits
    sink = LocalAttachmentSink(work_dir, limits)
    context = context.model_copy(update={"sink": sink, "limits": limits, "check_cancelled": check_cancelled or context.check_cancelled})
    setting = make_setting(path)
    documents = ExtractProcessor().extract(setting, context)
    result = ExtractionResult(
        documents=tuple(documents), attachments=tuple(asset.attachment for asset in sink.assets), source_sha256=hashlib.sha256(path.read_bytes()).hexdigest(), parse_fingerprint=canonical_parse_fingerprint(setting.profile)
    )
    assert decode_manifest(encode_manifest(result), limits) == result
    return documents, sink


def _png(path):
    with Image.new("RGB", (3, 2), "red") as image:
        image.save(path)


def _hyperlink(paragraph, text, url):
    element = OxmlElement("w:hyperlink")
    element.set(qn("r:id"), paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True))
    run = OxmlElement("w:r")
    child = OxmlElement("w:t")
    child.text = text
    run.append(child)
    element.append(run)
    paragraph._p.append(element)
    return element


def _image_pdf(path):
    """Three pages: same XObject twice on page 1, once on page 2, empty page 3."""
    write_pdf(path, ["first page", "", ""])
    writer = PdfWriter()
    writer.append(PdfReader(path))
    image = DecodedStreamObject()
    image.set_data(b"\xff\x00\x00" * 6)
    image.update(
        {
            NameObject("/Type"): NameObject("/XObject"),
            NameObject("/Subtype"): NameObject("/Image"),
            NameObject("/Width"): NumberObject(3),
            NameObject("/Height"): NumberObject(2),
            NameObject("/ColorSpace"): NameObject("/DeviceRGB"),
            NameObject("/BitsPerComponent"): NumberObject(8),
        }
    )
    reference = writer._add_object(image)
    for index, count in ((0, 2), (1, 1)):
        page = writer.pages[index]
        page["/Resources"][NameObject("/XObject")] = DictionaryObject({NameObject("/Im1"): reference})
        stream = DecodedStreamObject()
        stream.set_data(page.get_contents().get_data() + b"\n" + b"\n".join(f"q 30 0 0 20 {50 + occurrence * 50} 500 cm /Im1 Do Q".encode() for occurrence in range(count)))
        page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)
    writer.close()


def test_word_nested_table_order_repeats_and_run_spaces(tmp_path):
    path = tmp_path / "source.docx"
    document = WordFile()
    document.add_heading("设备", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("hello ")
    paragraph.add_run(" world").bold = True
    document.add_paragraph("重复文字")
    outer = document.add_table(rows=1, cols=2)
    outer.cell(0, 0).text = "外层"
    cell = outer.cell(0, 1)
    cell.text = "前"
    inner = cell.add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "嵌套"
    cell.add_paragraph("后")
    document.add_paragraph("重复文字")
    document.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    text = "\n".join(d.page_content for d in docs)
    assert "hello  world" in text and text.count("重复文字") == 2
    assert text.index("hello") < text.index("外层") < text.index("前") < text.index("嵌套") < text.index("后") < text.rindex("重复文字")
    assert docs[0].page_content == "# 设备"
    assert any(d.heading_path == ("设备",) for d in docs)
    assert any(s.location.get("table_path") == "1.1" and s.location["row"] == 1 for d in docs for s in d.source_spans)
    repeated = [s.location["paragraph"] for d in docs for s in d.source_spans if d.page_content == "重复文字"]
    assert repeated == [3, 4]


def test_word_merged_physical_cell_once_independent_equal_cells_repeat(tmp_path):
    path = tmp_path / "merged.docx"
    word = WordFile()
    table = word.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(1, 1)).text = "合并内容"
    table.cell(0, 2).text = "重复"
    table.cell(1, 2).text = "重复"
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    text = "\n".join(d.page_content for d in docs)
    assert text.count("合并内容") == 1
    assert text.count("重复") == 2
    assert all(d.kind == "table_row" for d in docs)
    assert "---" not in text  # No actual header evidence.


def test_word_explicit_repeat_header_has_header_source_spans(tmp_path):
    path = tmp_path / "header.docx"
    word = WordFile()
    table = word.add_table(rows=2, cols=2)
    for cell, value in zip(table.rows[0].cells, ("名称", "值")):
        cell.text = value
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    table.cell(1, 0).text = "设备"
    table.cell(1, 1).text = "A"
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    assert any("| 名称 | 值 |\n| --- | --- |" in d.page_content for d in docs)
    spans = [s for d in docs for s in d.source_spans]
    assert any(s.block_id.startswith("table_header:") and s.location["row"] == 1 and s.location["column"] == 2 for s in spans)


def test_word_hyperlinks_drawing_and_run_order_and_accurate_offsets(tmp_path):
    path = tmp_path / "images.docx"
    png = tmp_path / "red.png"
    _png(png)
    word = WordFile()
    word.add_heading("一级", level=1)
    word.add_heading("二级", level=2)
    p = word.add_paragraph()
    p.add_run("before ")
    _hyperlink(p, " docs ", "https://example.org/docs")
    _hyperlink(p, "unsafe", "javascript:alert(1)")
    run = p.add_run("left ")
    run.add_picture(str(png))
    run.add_text(" right")
    cell = word.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].add_run().add_picture(str(png))
    word.add_heading("新章", level=1)
    word.save(path)
    docs, sink = _extract(path, tmp_path / "work")
    assert len(sink.assets) == 1
    images = [(d, image) for d in docs for image in d.attachments]
    assert len(images) == 2
    assert images[0][0].heading_path == ("一级", "二级")
    text = images[0][0].page_content
    assert "[ docs ](https://example.org/docs)" in text
    assert "unsafe" in text and "javascript:" not in text
    assert text.index("left ") < images[0][1].source.start < text.index(" right")
    assert images[0][1].source.location == {"paragraph": 3, "image_index": 1}
    assert images[1][1].source.location["table"] == 1
    assert docs[-1].heading_path == ("新章",)
    for doc, occurrence in images:
        assert doc.page_content[occurrence.source.start : occurrence.source.end] == f"![图片](knowledge-attachment:{occurrence.ref})"
        assert occurrence.source in doc.source_spans


def test_word_legacy_field_hyperlink_spaces(tmp_path):
    path = tmp_path / "field.docx"
    word = WordFile()
    p = word.add_paragraph("before ")
    for kind in ("begin", "instruction", "separate", "text", "end"):
        run = p.add_run()
        if kind == "instruction":
            element = OxmlElement("w:instrText")
            element.text = ' HYPERLINK "https://example.org" '
        elif kind == "text":
            run.add_text(" field text ")
            continue
        else:
            element = OxmlElement("w:fldChar")
            element.set(qn("w:fldCharType"), kind)
        run._r.append(element)
    p.add_run(" after")
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    assert docs[0].page_content == "before [ field text ](https://example.org) after"


def test_word_remote_image_is_not_downloaded(tmp_path):
    path = tmp_path / "remote.docx"
    png = tmp_path / "red.png"
    _png(png)
    word = WordFile()
    p = word.add_paragraph("visible")
    p.add_run().add_picture(str(png))
    blip = p._p.xpath(".//a:blip")[0]
    blip.set(qn("r:embed"), word.part.relate_to("https://example.invalid/image.png", RT.IMAGE, is_external=True))
    word.save(path)
    docs, sink = _extract(path, tmp_path / "work")
    assert not sink.assets
    assert "visible" in docs[0].page_content
    assert any(w.code == "EXTERNAL_IMAGE_NOT_FETCHED" for d in docs for w in d.warnings)


def test_pdf_keeps_individual_pages_without_string_cache(tmp_path):
    path = tmp_path / "source.pdf"
    write_pdf(path, ["first page", "second page", ""])
    docs, _ = _extract(path, tmp_path / "work")
    assert len(docs) == 3
    assert [d.source_spans[0].location["page"] for d in docs] == [1, 2, 3]
    assert docs[0].page_content == "first page"
    assert docs[1].page_content == "second page"
    assert docs[2].page_content == ""


def test_pdf_images_keep_each_occurrence_with_page_only_sources(tmp_path):
    path = tmp_path / "image.pdf"
    _image_pdf(path)
    docs, sink = _extract(path, tmp_path / "work")
    assert len(docs) == 3 and len(sink.assets) == 1
    assert [len(d.attachments) for d in docs] == [2, 1, 0]
    assert [o.source.location for d in docs for o in d.attachments] == [{"page": 1, "image_index": 1}, {"page": 1, "image_index": 2}, {"page": 2, "image_index": 1}]
    assert docs[2].page_content == ""
    for doc in docs:
        for image in doc.attachments:
            assert image.alt_text == "本页图片"
            assert doc.page_content[image.source.start : image.source.end] == f"![本页图片](knowledge-attachment:{image.ref})"


@pytest.mark.parametrize("extension", [".pdf", ".docx"])
def test_cumulative_text_limit_is_enforced_inside_adapter(tmp_path, extension):
    path = tmp_path / ("budget" + extension)
    if extension == ".pdf":
        from actweave_knowledge.extraction.builtin.pdf_extractor import PdfExtractor

        adapter = PdfExtractor()
        write_pdf(path, ["123456", "abcdef"])
    else:
        from actweave_knowledge.extraction.builtin.word_extractor import WordExtractor

        adapter = WordExtractor()
        word = WordFile()
        word.add_paragraph("123456")
        word.add_paragraph("abcdef")
        word.save(path)
    context = make_context(tmp_path / "work").model_copy(update={"limits": ExtractionLimits(max_text_chars=10)})
    with pytest.raises(KnowledgeError) as error:
        adapter.extract(make_setting(path), context)
    assert error.value.code == KNOWLEDGE_QUOTA_EXCEEDED


def test_pdf_helper_escapes_literal_string_delimiters(tmp_path):
    path = tmp_path / "literal.pdf"
    write_pdf(path, [r"one (two) \ three"])
    assert PdfReader(path).pages[0].extract_text() == r"one (two) \ three"


def test_word_cell_paragraph_origins_and_repeats(tmp_path):
    path = tmp_path / "cell-paragraphs.docx"
    word = WordFile()
    cell = word.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "重复"
    cell.add_paragraph("重复")
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    assert docs[0].page_content.count("重复") == 2
    paragraphs = [s for s in docs[0].source_spans if ":paragraph:" in s.block_id]
    assert [s.location["paragraph"] for s in paragraphs] == [1, 2]
    assert len({s.block_id for s in paragraphs}) == 2
    for span in paragraphs:
        assert docs[0].page_content[span.start : span.end] == "重复"


def test_word_images_in_hyperlink_get_unique_indices_and_shifted_spans(tmp_path):
    path = tmp_path / "linked-image.docx"
    png = tmp_path / "red.png"
    _png(png)
    word = WordFile()
    p = word.add_paragraph("start ")
    p.add_run().add_picture(str(png))
    link = _hyperlink(p, "click ", "https://example.org")
    drawing = p.add_run()
    drawing.add_picture(str(png))
    link.append(drawing._r)
    p.add_run().add_picture(str(png))
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    assert [o.source.location["image_index"] for o in docs[0].attachments] == [1, 2, 3]
    assert len({o.source.block_id for o in docs[0].attachments}) == 3


def test_pdf_plaintext_markdown_cannot_hide_image_refs(tmp_path):
    path = tmp_path / "literal-image.pdf"
    _image_pdf(path)
    reader = PdfReader(path)
    writer = PdfWriter()
    writer.append(reader)
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 72 720 Td (~~~) Tj ET\nq 30 0 0 20 50 500 cm /Im1 Do Q")
    writer.pages[0][NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)
    writer.close()
    docs, _ = _extract(path, tmp_path / "work")
    assert len(docs[0].attachments) == 1


@pytest.mark.parametrize("extension", [".pdf", ".docx"])
def test_rejected_image_keeps_text_and_warning(tmp_path, extension):
    path = tmp_path / ("rejected" + extension)
    if extension == ".pdf":
        _image_pdf(path)
    else:
        png = tmp_path / "red.png"
        _png(png)
        word = WordFile()
        word.add_paragraph("visible").add_run().add_picture(str(png))
        word.save(path)
    docs, sink = _extract(path, tmp_path / "work", limits=ExtractionLimits(max_image_pixels=1))
    assert not sink.assets
    assert "visible" in docs[0].page_content if extension == ".docx" else "first page" in docs[0].page_content
    assert any(w.code == "IMAGE_LIMIT_EXCEEDED" for d in docs for w in d.warnings)


def test_pdf_native_handles_close_when_sink_fails(tmp_path, monkeypatch):
    import pypdfium2
    from actweave_knowledge.extraction.builtin.pdf_extractor import PdfExtractor

    path = tmp_path / "close.pdf"
    _image_pdf(path)
    closed = []
    for cls in (pypdfium2.PdfDocument, pypdfium2.PdfPage, pypdfium2.PdfTextPage, pypdfium2.PdfImage, pypdfium2.PdfBitmap):
        original = cls.close

        def record_close(self, _original=original, _name=cls.__name__, **kwargs):
            closed.append(_name)
            return _original(self, **kwargs)

        monkeypatch.setattr(cls, "close", record_close)

    class FailingSink:
        def accept(self, *args, **kwargs):
            raise OSError("sink failed")

    context = make_context(tmp_path / "work").model_copy(update={"sink": FailingSink()})
    with pytest.raises(OSError, match="sink failed"):
        PdfExtractor().extract(make_setting(path), context)
    assert {"PdfDocument", "PdfPage", "PdfTextPage", "PdfImage", "PdfBitmap"} <= set(closed)
    assert not list(context.work_dir.iterdir())


def test_pdf_native_handles_close_on_text_budget(tmp_path, monkeypatch):
    import pypdfium2
    from actweave_knowledge.extraction.builtin.pdf_extractor import PdfExtractor

    path = tmp_path / "close-text.pdf"
    write_pdf(path, ["123456", "not visited"])
    closed = []
    for cls in (pypdfium2.PdfDocument, pypdfium2.PdfPage, pypdfium2.PdfTextPage):
        original = cls.close

        def record_close(self, _original=original, _name=cls.__name__, **kwargs):
            closed.append(_name)
            return _original(self, **kwargs)

        monkeypatch.setattr(cls, "close", record_close)
    context = make_context(tmp_path / "work").model_copy(update={"limits": ExtractionLimits(max_text_chars=5)})
    with pytest.raises(KnowledgeError):
        PdfExtractor().extract(make_setting(path), context)
    assert {"PdfDocument", "PdfPage", "PdfTextPage"} <= set(closed)


@pytest.mark.parametrize("extension", [".pdf", ".docx"])
def test_adapter_propagates_cancellation_during_extraction(tmp_path, extension):
    path = tmp_path / ("cancel" + extension)
    if extension == ".pdf":
        from actweave_knowledge.extraction.builtin.pdf_extractor import PdfExtractor

        adapter = PdfExtractor()
        _image_pdf(path)
    else:
        from actweave_knowledge.extraction.builtin.word_extractor import WordExtractor

        adapter = WordExtractor()
        word = WordFile()
        word.add_paragraph("before")
        word.add_paragraph("after")
        word.save(path)
    calls = 0

    class Cancelled(Exception):
        pass

    def guard():
        nonlocal calls
        calls += 1
        if calls == 4:
            raise Cancelled()

    context = make_context(tmp_path / "work").model_copy(update={"check_cancelled": guard})
    with pytest.raises(Cancelled):
        adapter.extract(make_setting(path), context)
    assert calls == 4


def test_pdf_soft_mask_survives_safe_attachment_normalization(tmp_path):
    path = tmp_path / "transparent.pdf"
    _image_pdf(path)
    writer = PdfWriter()
    writer.append(PdfReader(path))
    mask = DecodedStreamObject()
    mask.set_data(bytes([0, 128, 255, 0, 128, 255]))
    mask.update(
        {
            NameObject("/Type"): NameObject("/XObject"),
            NameObject("/Subtype"): NameObject("/Image"),
            NameObject("/Width"): NumberObject(3),
            NameObject("/Height"): NumberObject(2),
            NameObject("/ColorSpace"): NameObject("/DeviceGray"),
            NameObject("/BitsPerComponent"): NumberObject(8),
        }
    )
    image = writer.pages[0]["/Resources"]["/XObject"]["/Im1"]
    image[NameObject("/SMask")] = writer._add_object(mask)
    with path.open("wb") as output:
        writer.write(output)
    writer.close()
    docs, sink = _extract(path, tmp_path / "work")
    assert [len(d.attachments) for d in docs] == [2, 1, 0]
    with Image.open(tmp_path / "work" / sink.assets[0].relative_path) as normalized:
        assert normalized.getchannel("A").getextrema() == (0, 255)
        with normalized.getchannel("A") as alpha:
            assert 128 in alpha.tobytes()


def test_word_leading_run_spaces_do_not_turn_image_into_code(tmp_path):
    path = tmp_path / "indent.docx"
    png = tmp_path / "red.png"
    _png(png)
    word = WordFile()
    p = word.add_paragraph()
    p.add_run("  ")
    p.add_run("  visible ")
    p.add_run().add_picture(str(png))
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    assert len(docs[0].attachments) == 1
    from markdown_it import MarkdownIt

    rendered = MarkdownIt().render(docs[0].page_content)
    assert "code>" not in rendered
    assert "   visible " in rendered


def test_word_rejected_image_retains_actual_later_image_index(tmp_path):
    path = tmp_path / "indices.docx"
    large, small = tmp_path / "large.png", tmp_path / "small.png"
    _png(large)
    with Image.new("RGB", (1, 1), "blue") as image:
        image.save(small)
    word = WordFile()
    p = word.add_paragraph("text")
    p.add_run().add_picture(str(large))
    p.add_run().add_picture(str(small))
    word.save(path)
    docs, sink = _extract(path, tmp_path / "work", limits=ExtractionLimits(max_image_pixels=1))
    assert len(sink.assets) == 1
    assert docs[0].attachments[0].source.location["image_index"] == 2
    assert docs[0].warnings[0].source_position["image_index"] == 1


def test_word_two_paragraphs_keep_only_their_own_source_intervals(tmp_path):
    path = tmp_path / "paragraphs.docx"
    word = WordFile()
    word.add_paragraph("first")
    word.add_paragraph("second paragraph split target")
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    assert len(docs) == 2
    assert [(s.start, s.end, s.location) for d in docs for s in d.source_spans] == [(0, 5, {"paragraph": 1}), (0, 29, {"paragraph": 2})]


@pytest.mark.parametrize("extension", [".pdf", ".docx"])
def test_image_intermediates_respect_work_directory_limit(tmp_path, extension):
    path = tmp_path / ("work-limit" + extension)
    if extension == ".pdf":
        from actweave_knowledge.extraction.builtin.pdf_extractor import PdfExtractor

        adapter = PdfExtractor()
        _image_pdf(path)
    else:
        from actweave_knowledge.extraction.builtin.word_extractor import WordExtractor

        adapter = WordExtractor()
        png = tmp_path / "red.png"
        _png(png)
        word = WordFile()
        word.add_paragraph("image").add_run().add_picture(str(png))
        word.save(path)
    from actweave_knowledge.extraction.contracts import ExtractionError

    context = make_context(tmp_path / "work").model_copy(update={"limits": ExtractionLimits(max_work_dir_bytes=10)})
    with pytest.raises(ExtractionError) as error:
        adapter.extract(make_setting(path), context)
    assert error.value.reason_code == "PARSER_WORK_DIR_LIMIT_EXCEEDED"
    assert not list(context.work_dir.iterdir())


def test_word_explicit_header_merged_cells_keep_markdown_columns(tmp_path):
    path = tmp_path / "merged-header.docx"
    word = WordFile()
    table = word.add_table(rows=2, cols=3)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    table.cell(0, 0).merge(table.cell(0, 1)).text = "merged"
    table.cell(0, 2).text = "last"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "B"
    table.cell(1, 2).text = "C"
    word.save(path)
    docs, _ = _extract(path, tmp_path / "work")
    assert docs[0].page_content == "| merged |  | last |\n| --- | --- | --- |"
    assert docs[1].page_content == "| A | B | C |"


@pytest.mark.parametrize("flate", [False, True], ids=["uncompressed", "flate"])
def test_pdf_cmyk_image_uses_rendered_fallback_without_losing_page_text(tmp_path, flate):
    path = tmp_path / "cmyk.pdf"
    write_pdf(path, ["visible CMYK page"])
    writer = PdfWriter()
    writer.append(PdfReader(path))
    image = DecodedStreamObject()
    image.set_data(bytes([0, 255, 255, 0]) * 6)
    image.update(
        {
            NameObject("/Type"): NameObject("/XObject"),
            NameObject("/Subtype"): NameObject("/Image"),
            NameObject("/Width"): NumberObject(3),
            NameObject("/Height"): NumberObject(2),
            NameObject("/ColorSpace"): NameObject("/DeviceCMYK"),
            NameObject("/BitsPerComponent"): NumberObject(8),
        }
    )
    reference = writer._add_object(image.flate_encode() if flate else image)
    page = writer.pages[0]
    page["/Resources"][NameObject("/XObject")] = DictionaryObject({NameObject("/CMYK"): reference})
    stream = DecodedStreamObject()
    stream.set_data(page.get_contents().get_data() + b"\nq 30 0 0 20 50 500 cm /CMYK Do Q")
    page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)
    writer.close()

    docs, sink = _extract(path, tmp_path / "work")
    assert len(docs) == 1 and "visible CMYK page" in docs[0].page_content
    assert len(docs[0].attachments) == 1 and len(sink.assets) == 1
    assert docs[0].attachments[0].source.location == {"page": 1, "image_index": 1}
    assert not docs[0].warnings
    with Image.open(tmp_path / "work" / sink.assets[0].relative_path) as normalized:
        assert normalized.mode == "RGBA" and normalized.size == (3, 2)
        red, green, blue, alpha = normalized.getpixel((0, 0))
        assert red > green and red > blue and alpha == 255


@pytest.mark.parametrize("break_tag", ["w:br", "w:cr"])
def test_word_multiline_header_and_cells_preserve_one_markdown_row_and_image_offsets(tmp_path, break_tag):
    from markdown_it import MarkdownIt

    path = tmp_path / "multiline-table.docx"
    png = tmp_path / "red.png"
    _png(png)
    word = WordFile()
    word.add_paragraph("outside\nparagraph")
    table = word.add_table(rows=2, cols=2)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    header = table.cell(0, 0).paragraphs[0].add_run("Name")
    header._r.append(OxmlElement(break_tag))
    header.add_text("caption")
    table.cell(0, 1).text = "Value"
    run = table.cell(1, 0).paragraphs[0].add_run("line1")
    run._r.append(OxmlElement(break_tag))
    run.add_text("line2 ")
    run.add_picture(str(png))
    table.cell(1, 1).text = "42"
    word.save(path)

    docs, _ = _extract(path, tmp_path / "work")
    assert docs[0].page_content == "outside\nparagraph"
    header_doc, row_doc = docs[1:]
    assert header_doc.page_content == "| Name; caption | Value |\n| --- | --- |"
    assert row_doc.page_content.startswith("| line1; line2 ![图片]")
    assert row_doc.page_content.endswith(" | 42 |")
    assert "\n" not in row_doc.page_content and "<br" not in row_doc.page_content
    combined = header_doc.page_content + "\n" + row_doc.page_content
    tokens = MarkdownIt().enable("table").parse(combined)
    assert sum(token.type == "table_open" for token in tokens) == 1
    assert sum(token.type == "tr_open" for token in tokens) == 2
    assert sum(token.type == "td_open" for token in tokens) == 2
    assert sum(token.type == "th_open" for token in tokens) == 2
    assert combined.count("line1") == 1 and combined.count("line2") == 1
    header_span = next(span for span in header_doc.source_spans if span.block_id == "table_header:1:column:1")
    assert header_doc.page_content[header_span.start : header_span.end] == "Name; caption"
    occurrence = row_doc.attachments[0]
    assert row_doc.page_content[occurrence.source.start : occurrence.source.end] == f"![图片](knowledge-attachment:{occurrence.ref})"
    assert occurrence.source.location == {"table": 1, "table_path": "1", "row": 2, "column": 1, "paragraph": 1, "image_index": 1}
    cell_span = next(span for span in row_doc.source_spans if span.block_id == "table:1:row:2:column:1")
    paragraph_span = next(span for span in row_doc.source_spans if span.block_id == "table:1:row:2:column:1:paragraph:1")
    assert (cell_span.start, cell_span.end) == (paragraph_span.start, paragraph_span.end)
    assert row_doc.page_content[cell_span.start : cell_span.end] == f"line1; line2 ![图片](knowledge-attachment:{occurrence.ref})"

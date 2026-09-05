"""Structured chunking gates: budgets, source intersections, and legacy units."""

from __future__ import annotations

import pytest
from actweave_knowledge.contracts import KNOWLEDGE_MAX_SEGMENT_CHARS
from actweave_knowledge.extraction.contracts import AttachmentOccurrence, Document, ExtractionError, SourceSpan
from actweave_knowledge.ingestion import splitter
from actweave_knowledge.ingestion.structure import StructureUnit, inline_atoms, join_units
from actweave_knowledge.ingestion.tokenizer import count_knowledge_tokens
from parsing_test_helpers import make_chunk_profile, make_document


def split(documents, warnings=None, **overrides):
    return splitter.split_documents(tuple(documents), profile=make_chunk_profile(**overrides), warnings=warnings)


@pytest.mark.parametrize("field", ["splitter_version", "cleaner_version"])
def test_token_splitter_rejects_unavailable_version(field):
    profile = make_chunk_profile(**{field: "old-unavailable-build"})
    with pytest.raises(ExtractionError) as error:
        splitter.split_documents((make_document("alpha beta"),), profile=profile)
    assert error.value.reason_code == "PROCESSING_PROFILE_UNAVAILABLE"


@pytest.mark.parametrize("text,end", [(r"\# literal", 2), ("&#32;x", 5), ("&#9;x", 4), ("&amp;x", 5)])
def test_literal_escape_and_entity_are_atomic(text, end):
    assert (0, end) in inline_atoms(text, include_text_escapes=True)


def test_character_fallback_cannot_split_a_markdown_escape():
    with pytest.raises(ExtractionError) as error:
        list(splitter._split_ranges(StructureUnit(r"\#"), [""], lambda part: len(part.content) <= 1))
    assert error.value.reason_code == "ATOMIC_CONTENT_EXCEEDS_BUDGET"


def test_split_atom_protection_does_not_disable_plain_url_cleaning():
    from actweave_knowledge.ingestion.cleaner import clean_documents

    document = make_document(r"before https://example.test/a\_b after")
    [cleaned] = clean_documents((document,), remove_extra_spaces=False, remove_urls_emails=True)
    assert "https://" not in cleaned.page_content
    assert inline_atoms(r"\# plain") == []
    assert inline_atoms(r"\# plain", include_text_escapes=True) == [(0, 2)]


@pytest.mark.parametrize(
    "text,expected",
    [
        (r"\`alice@example.test\`", r"\`\`"),
        (r"\[alice@example.test\](/literal)", r"\[\](/literal)"),
        (r"\!\[alice@example.test\](/literal)", r"\!\[\](/literal)"),
    ],
)
def test_cleaner_consumes_escaped_fake_markdown_before_email_removal(text, expected):
    from actweave_knowledge.ingestion.cleaner import clean_documents

    [cleaned] = clean_documents((make_document(text),), remove_extra_spaces=False, remove_urls_emails=True)
    assert cleaned.page_content == expected


def test_cleaner_preserves_real_markdown_and_respects_backslash_parity():
    from actweave_knowledge.ingestion.cleaner import clean_documents

    text = r"`alice@example.test` [alice@example.test](/literal) \\`alice@example.test` \\\`alice@example.test\`"
    [cleaned] = clean_documents((make_document(text),), remove_extra_spaces=False, remove_urls_emails=True)
    assert "`alice@example.test`" in cleaned.page_content
    assert "[alice@example.test](/literal)" in cleaned.page_content
    assert r"\\`alice@example.test`" in cleaned.page_content
    assert r"\\\`\`" in cleaned.page_content


def assert_budgets(drafts, size=200, child_size=100):
    assert drafts
    assert [d.position for d in drafts] == list(range(1, len(drafts) + 1))
    for draft in drafts:
        for value, limit in [(draft, size), *((child, child_size) for child in draft.children)]:
            assert len(value.content) <= KNOWLEDGE_MAX_SEGMENT_CHARS
            assert count_knowledge_tokens(value.content) <= limit
            assert count_knowledge_tokens(value.index_text) <= limit
            assert value.token_count == count_knowledge_tokens(value.index_text)
            assert value.index_text.strip()
            assert all(0 <= s.start < s.end <= len(value.content) for s in value.source_spans)


@pytest.mark.parametrize(
    ("kind", "text"),
    [
        ("paragraph", " ".join(f"word{i}" for i in range(350))),
        ("paragraph", "\n".join(f"line{i} " + "plain " * 13 for i in range(40))),
        ("paragraph", "".join(chr(0x4E00 + i) for i in range(500))),
        ("page", "\n\n".join((f"para{i} " + "normal " * 45).strip() for i in range(5))),
        ("text", "\n\n".join((f"para{i} " + "normal " * 45).strip() for i in range(5))),
    ],
)
def test_body_overlap_consumes_new_source_without_inserting_separators(kind, text):
    document = make_document(text, location={"page": 1}).model_copy(update={"kind": kind})
    zero = split([document], size=200, overlap=0, child_size=100)
    drafts = split([document], size=200, overlap=60, child_size=100)
    assert [draft.content for draft in drafts] != [draft.content for draft in zero]
    assert_budgets(drafts)
    previous_start, previous_end, duplicated = -1, 0, False
    for draft in drafts:
        start = text.find(draft.content, previous_start + 1)
        assert start >= 0
        end = start + len(draft.content)
        assert end > previous_end
        assert not text[previous_end:start].strip()
        if start < previous_end:
            duplicated = True
            shared = text[start:previous_end]
            assert count_knowledge_tokens(shared) <= 60
            assert count_knowledge_tokens(splitter.build_index_text(shared)) <= 60
        assert all(span.role == "source" and span.location == {"page": 1} for span in draft.source_spans)
        previous_start, previous_end = start, end
    assert duplicated
    assert not text[previous_end:].strip()


def test_overlap_suffix_uses_user_boundary_before_sentence_fallback():
    text = "前" * 200 + "分界" + "章" * 10 + "。" + "节" * 4
    unit = StructureUnit(text, (SourceSpan(block_id="body", start=0, end=len(text), location={"page": 2}),), kind="page")
    tail = splitter._overlap_tail([unit], "\n\n", 50, ["分界", "。", ""])
    assert len(tail) == 1
    assert tail[0].content == "章" * 10 + "。" + "节" * 4
    assert tail[0].source_spans[0].role == "source"
    assert tail[0].source_spans[0].location == {"page": 2}
    assert splitter.fits_chunk(tail[0].content, 50)


@pytest.mark.parametrize("marker", ["- item", "1. item"])
def test_overlap_suffix_preserves_inline_block_marker_and_source(marker):
    text = "alpha " * 80 + "boundary " + marker
    unit = StructureUnit(text, (SourceSpan(block_id="body", start=0, end=len(text), location={"page": 2}),))
    [tail] = splitter._overlap_tail([unit], "\n\n", 10, ["boundary ", " ", ""])
    assert splitter.build_index_text(tail.content) == marker
    assert tail.kind == "paragraph"
    assert "".join(tail.content[s.start : s.end] for s in tail.source_spans) == marker
    assert all(s.role == "source" and s.location == {"page": 2} for s in tail.source_spans)
    assert splitter.fits_chunk(tail.content, 10)
    assert not tail.attachments


@pytest.mark.parametrize("marker", ["- item", "1. item"])
@pytest.mark.parametrize("size,leading", [(200, 198), (400, 98)])
def test_split_boundaries_preserve_inline_block_markers_in_parents_and_children(marker, size, leading):
    from markdown_it import MarkdownIt

    text = "alpha " * leading + "boundary " + marker + " " + "omega " * 250
    drafts = split([make_document(text)], mode="parent_child", size=size, overlap=0, child_size=100, separator="boundary ", child_separator="boundary ")
    assert_budgets(drafts, size=size)
    assert sum(len(draft.children) for draft in drafts) > 1
    for values in (drafts, [child for draft in drafts for child in draft.children]):
        assert "".join(value.index_text for value in values).replace(" ", "") == text.replace(" ", "")
        source = "".join(value.content[s.start : s.end] for value in values for s in value.source_spans if s.role == "source")
        assert source.replace(" ", "") == text.replace(" ", "")
        assert all(s.location == {"paragraph": 1} for value in values for s in value.source_spans)
        assert not any(token.type in {"bullet_list_open", "ordered_list_open"} for value in values for token in MarkdownIt("commonmark").parse(value.content))


@pytest.mark.parametrize("indented", [False, True])
def test_split_boundary_protection_does_not_escape_native_code_body(indented):
    body = "- item\n1. item\n" * 150
    text = "".join("    " + line for line in body.splitlines(keepends=True)) if indented else "```\n" + body + "```"
    drafts = split([make_document(text)], mode="parent_child", size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    for values in (drafts, [child for draft in drafts for child in draft.children]):
        assert "\n".join(value.index_text for value in values) == body.rstrip()
        assert all(value.content.startswith("```") for value in values)


@pytest.mark.parametrize("text,block_type,indexed", [("- one\n- two", "bullet_list_open", "one\ntwo"), ("| key |\n| --- |\n| value |", "table_open", "key\nvalue")])
def test_split_boundary_protection_preserves_native_list_and_table(text, block_type, indexed):
    from markdown_it import MarkdownIt

    [draft] = split([make_document(text)], mode="parent_child", size=200, overlap=0, child_size=100)
    for value in (draft, *draft.children):
        assert value.content == text
        assert value.index_text == indexed
        assert any(token.type == block_type for token in MarkdownIt("commonmark").enable("table").parse(value.content))


def test_split_boundary_protection_shifts_attachment_and_preserves_source():
    text = "alpha " * 198 + "boundary - item ![image](knowledge-attachment:" + "a" * 64 + ") " + "omega " * 100
    start = text.index("![image]")
    end = text.index(")", start) + 1
    occurrence = AttachmentOccurrence(ref="a" * 64, alt_text="image", source=SourceSpan(block_id="body", start=start, end=end, location={"page": 1}))
    document = make_document(text, location={"page": 1}).model_copy(update={"attachments": (occurrence,)})
    drafts = split([document], mode="parent_child", size=200, overlap=0, child_size=100, separator="boundary ")
    assert_budgets(drafts)
    [(draft, image)] = [(draft, image) for draft in drafts for image in draft.attachments]
    assert draft.content[image.source.start : image.source.end] == text[start:end]
    assert image.source.location == occurrence.source.location
    assert image.ref == occurrence.ref
    assert draft.index_text.startswith("- item image")
    source = "".join(draft.content[s.start : s.end] for draft in drafts for s in draft.source_spans)
    assert source.replace(" ", "") == text.replace(" ", "")


@pytest.mark.parametrize("indent,protected", [("    ", "&#32;   "), ("\t", "&#9;")])
def test_fragment_indentation_protection_remaps_source_and_attachment(indent, protected):
    from actweave_knowledge.ingestion.structure import slice_fragment

    image = "![image](knowledge-attachment:" + "a" * 64 + ")"
    text = "before" + indent + r"\# literal \&amp; " + image
    image_start = text.index("![image]")
    occurrence = AttachmentOccurrence(ref="a" * 64, alt_text="image", source=SourceSpan(block_id="body", start=image_start, end=len(text), location={"page": 1}))
    unit = StructureUnit(text, (SourceSpan(block_id="body", start=0, end=len(text), location={"page": 1}),), attachments=(occurrence,))
    fragment = slice_fragment(unit, len("before"), len(text))
    assert fragment.content == protected + r"\# literal \&amp; " + image
    assert splitter.build_index_text(fragment.content) == "# literal &amp; image"
    [attached] = fragment.attachments
    assert fragment.content[attached.source.start : attached.source.end] == image
    assert attached.source.location == occurrence.source.location and attached.ref == occurrence.ref
    assert "".join(fragment.content[s.start : s.end] for s in fragment.source_spans) == fragment.content
    assert all(s.role == "source" and s.location == {"page": 1} for s in fragment.source_spans)
    assert fragment.kind == "paragraph"
    assert splitter.fits_chunk(fragment.content, 100)


def test_overlap_keeps_whole_lists_but_does_not_slice_list_fragments():
    text = "- one\n- two"
    unit = StructureUnit(text, (SourceSpan(block_id="list", start=0, end=len(text), location={"paragraph": 1}),), kind="list")
    assert splitter._overlap_tail([unit], "\n\n", 50, ["\n", " ", ""]) == [unit]
    assert splitter._overlap_tail([unit], "\n\n", 1, ["\n", " ", ""]) == []
    fragment = splitter.replace(unit, kind="list_fragment")
    assert splitter._overlap_tail([fragment], "\n\n", 50, ["\n", " ", ""]) == []


def test_full_incoming_atomic_prefix_shrinks_tail_without_splitting_link():
    old_text = "旧" * 30
    old = StructureUnit(old_text, (SourceSpan(block_id="old", start=0, end=len(old_text), location={"paragraph": 1}),), kind="text_fragment")
    link = "[link](https://example.test/" + "x/" * 150 + ")"
    text = link + "新" * 100
    incoming = StructureUnit(text, (SourceSpan(block_id="new", start=0, end=len(text), location={"paragraph": 2}),), kind="text_fragment")
    assert splitter.fits_chunk(link, 200)
    assert not splitter.fits_chunk(old_text + link, 200)
    retained, first, rest = splitter._reserve_overlap(
        StructureUnit(""),
        [old],
        incoming,
        "\n\n",
        limit=200,
        overlap=60,
        separators=["\n\n", " ", ""],
        continuation=True,
    )
    assert retained and len(retained[0].content) < len(old_text)
    assert first.content.startswith(link)
    assert rest is not None
    assert first.content + rest.content == text
    assert splitter.fits_chunk(join_units(retained, "\n\n").content, 60)
    packed = splitter._append_piece(retained, first, continuation=True)
    assert splitter.fits_chunk(join_units(packed, "\n\n").content, 200)
    assert all(span.role == "source" for value in [*retained, first, rest] for span in value.source_spans)
    assert inline_atoms(first.content, include_text_escapes=True)[0] == (0, len(link))


def test_positive_overlap_with_nearly_full_budget_still_terminates():
    text = "".join(chr(0x5000 + i) for i in range(220))
    drafts = split([make_document(text)], size=200, overlap=190, child_size=100)
    assert 1 < len(drafts) <= len(text)
    assert_budgets(drafts)
    previous_start, previous_end, duplicated = -1, 0, False
    for draft in drafts:
        start = text.find(draft.content, previous_start + 1)
        assert start >= 0
        end = start + len(draft.content)
        assert end > previous_end
        duplicated = duplicated or start < previous_end
        previous_start, previous_end = start, end
    assert previous_end == len(text)
    assert duplicated


@pytest.mark.parametrize("encoded", [r"\#", "&#32;", "&#9;"])
def test_overlap_never_cuts_literal_serialization_atoms(encoded):
    text = "".join(chr(0x4E00 + index) + encoded for index in range(180))
    atoms = inline_atoms(text, include_text_escapes=True)
    drafts = split([make_document(text)], size=200, overlap=60, child_size=100)
    assert_budgets(drafts)
    previous_start, previous_end = -1, 0
    for draft in drafts:
        start = text.find(draft.content, previous_start + 1)
        assert start >= 0
        end = start + len(draft.content)
        assert not any(a < start < b or a < end < b for a, b in atoms)
        assert end > previous_end
        previous_start, previous_end = start, end
    assert previous_end == len(text)


def test_table_continuations_keep_header_and_source():
    rows = ["| 编号 | 处置 |", "| --- | --- |"] + [f"| E{i:03d} | 检查邻居并核对接口状态。" + "确认链路。" * 12 + " |" for i in range(20)]
    text = "\n".join(rows)
    spans = []
    offset = 0
    for row, line in enumerate(rows, 1):
        spans.append(SourceSpan(block_id=f"row:{row}", start=offset, end=offset + len(line), location={"row": row}))
        offset += len(line) + 1
    drafts = split([Document(page_content=text, source_spans=tuple(spans))], size=200, overlap=0, child_size=100)
    assert len(drafts) > 1
    assert_budgets(drafts)
    assert all("编号" in d.content and "处置" in d.content for d in drafts)
    assert sum(d.content.count(f"E{i:03d}") for d in drafts for i in range(20)) == 20
    assert all(s.role == "context_prefix" for d in drafts[1:] for s in d.source_spans if s.location["row"] <= 2)


def test_long_code_line_splits_unicode_without_losing_payload():
    value = "告警🧑🏽‍💻List<int>" * 100
    drafts = split([make_document("```cpp\n" + value + "\n```")], size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    assert "".join(d.content.removeprefix("```cpp\n").removesuffix("\n```") for d in drafts) == value


def test_source_intersections_after_merging_and_mid_paragraph_split():
    texts = ["A " * 40, "B " * 270, "C " * 40]
    docs = [make_document(t.strip(), location={"paragraph": i}) for i, t in enumerate(texts, 1)]
    drafts = split(docs, size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    assert {s.location["paragraph"] for s in drafts[0].source_spans} == {1, 2}
    assert {s.location["paragraph"] for s in drafts[-1].source_spans} == {2, 3}
    assert all(d.content[s.start : s.end].strip().startswith({1: "A", 2: "B", 3: "C"}[s.location["paragraph"]]) for d in drafts for s in d.source_spans if d.content[s.start : s.end].strip())


@pytest.mark.parametrize("kind", ["heading", "table"])
def test_required_prefix_over_budget_degrades_with_a_warning_instead_of_failing(kind):
    """A heading path or table header wider than the chunk budget used to fail
    the whole document. It now truncates the context prefix, keeps every body
    row/paragraph, and reports one safe warning."""

    text = "# " + "标题" * 500 + "\n\n正文。" if kind == "heading" else "| " + "列名" * 500 + " |\n| --- |\n| 数据 |"
    warnings: list = []
    drafts = split([make_document(text)], warnings, size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    body = [d for d in drafts if ("正文" if kind == "heading" else "数据") in d.content]
    assert len(body) == 1
    # The body chunk keeps a shortened context prefix and at least half the
    # budget for its own text; the heading/header text itself is not lost —
    # it was emitted in full as ordinary chunks before the body.
    assert count_knowledge_tokens(body[0].content) <= 200
    marker = "标题" if kind == "heading" else "列名"
    assert sum(d.content.count(marker) for d in drafts if d is not body[0]) == 500
    assert [warning.code for warning in warnings] == ["OVERSIZED_PREFIX_SPLIT", "CONTEXT_PREFIX_TRUNCATED"]
    assert warnings[0].source_position == {"paragraph": 1}


def test_deep_heading_path_drops_top_levels_before_truncating():
    """The leaf heading is the most specific context: outer levels go first
    and the leaf survives intact when that alone fits."""

    text = "# " + "总纲" * 120 + "\n\n## " + "章节" * 120 + "\n\n### 叶子标题\n\n" + "正文内容。" * 20
    warnings: list = []
    drafts = split([make_document(text)], warnings, size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    body = [d for d in drafts if "正文内容" in d.content]
    assert body and all(d.content.startswith("### 叶子标题") for d in body)
    assert all("总纲" not in d.content and "章节" not in d.content for d in body)
    # Every heading's own text survives somewhere as source text.
    source = "".join(d.content[s.start : s.end] for d in drafts for s in d.source_spans if s.role == "source")
    assert source.count("总纲") == 120 and source.count("章节") == 120 and "叶子标题" in source
    assert "CONTEXT_PREFIX_TRUNCATED" in [warning.code for warning in warnings]


def test_display_budget_and_character_cap_are_independent_of_index_budget():
    text = "[可见](https://example.test/" + "x" * 800 + ")\n\n" + "a " * 6000
    drafts = split([make_document(text)], size=4000, overlap=0, child_size=500)
    assert_budgets(drafts, size=4000, child_size=500)
    assert sum(d.content.count("https://example.test/") for d in drafts) == 1


def test_word_level_packing_re_encodes_only_the_appended_words(monkeypatch: pytest.MonkeyPatch):
    """Measuring the whole candidate per word makes an unpunctuated paragraph quadratic in its length."""
    from actweave_knowledge.ingestion import tokenizer

    real = tokenizer._encoder()
    encoded_lengths: list[int] = []

    class RecordingEncoder:
        def encode_ordinary(self, value: str) -> list[int]:
            encoded_lengths.append(len(value))
            return real.encode_ordinary(value)

    monkeypatch.setattr(tokenizer, "_encoder", RecordingEncoder)
    text = "# 无标点长段\n\n" + " ".join(f"word{i}" for i in range(4000))
    drafts = split([make_document(text)], size=1000, overlap=0)
    encoded_during_split = sum(encoded_lengths)
    assert_budgets(drafts, size=1000)
    assert "".join(d.content for d in drafts).count("word") == 4000
    assert encoded_during_split < 40 * len(text)


def test_large_token_budgets_are_reachable_for_english_text():
    """4000 tokens of English is ~16000 characters: the old 4000-character
    ceiling silently capped every English chunk near 1000 tokens."""

    text = "word " * 3000  # ~3000 tokens, 15000 characters
    drafts = split([make_document(text.strip())], size=4000, overlap=0, child_size=500)
    assert_budgets(drafts, size=4000, child_size=500)
    assert len(drafts) == 1
    assert len(drafts[0].content) > 4000


def test_chinese_sentence_punctuation_is_a_split_boundary():
    """Sentences ending in ！/？/； no longer fall through to space or
    character cuts: every chunk ends on a sentence boundary."""

    sentences = [f"第{i}句提醒请检查设备状态是否正常！" for i in range(40)]
    text = "".join(sentences)
    drafts = split([make_document(text)], size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    assert len(drafts) > 1
    assert all(d.content.endswith("！") for d in drafts)
    assert "".join(d.content for d in drafts) == text
    questions = "".join(f"第{i}个问题是什么？" for i in range(60))
    question_drafts = split([make_document(questions)], size=200, overlap=0, child_size=100)
    assert len(question_drafts) > 1 and all(d.content.endswith("？") for d in question_drafts)


def test_pdf_pages_and_tabular_data_rows_do_not_merge_or_overlap():
    docs = [make_document("第一页。" * 100, location={"page": 1}), make_document("第二页。" * 100, location={"page": 2})]
    drafts = split(docs, size=200, overlap=100, child_size=100)
    assert all(len({s.location["page"] for s in d.source_spans}) == 1 for d in drafts)
    rows = [make_document(f"- 标识: {i}\n- 值: " + "value " * 80, location={"sheet": "S", "row": i}).model_copy(update={"kind": "table_row"}) for i in (1, 2)]
    row_drafts = split(rows, size=200, overlap=100, child_size=100)
    assert len(row_drafts) == 2
    assert [d.source_position["row"] for d in row_drafts] == [1, 2]


def test_children_use_own_separator_and_sources_without_repeating_images():
    ref = "a" * 64
    text = "第一部分。" * 30 + "分界" + "第二部分。" * 30 + f"\n\n![拓扑](knowledge-attachment://{ref})"
    image_start = text.index("![")
    occurrence = AttachmentOccurrence(ref=ref, alt_text="拓扑", source=SourceSpan(block_id="image", start=image_start, end=len(text), location={"paragraph": 1}))
    doc = make_document(text).model_copy(update={"attachments": (occurrence,)})
    drafts = split([doc], size=1000, overlap=0, child_size=100, mode="parent_child", child_separator="分界")
    assert_budgets(drafts, size=1000)
    assert len(drafts) == 1 and len(drafts[0].children) > 1
    assert drafts[0].attachments == (occurrence,)
    assert sum(c.content.count("第一部分。") for c in drafts[0].children) == 30
    assert sum(c.content.count("第二部分。") for c in drafts[0].children) == 30


def test_pure_image_source_does_not_create_indexable_segment():
    assert split([make_document(f"![不作为正文](knowledge-attachment://{'a' * 64})")]) == []


def test_cleaning_protects_code_and_internal_refs_and_remaps_sources():
    from actweave_knowledge.ingestion import cleaner

    ref = "a" * 64
    text = f"联系 a@example.test  后续。\n\n`a@example.test  x`\n\n```text\nhttps://example.test  x\n```\n\n![图](knowledge-attachment://{ref})"
    cleaned = cleaner.clean_documents((make_document(text),), remove_extra_spaces=True, remove_urls_emails=True)[0]
    assert cleaned.page_content.startswith("联系 后续。")
    assert "`a@example.test  x`" in cleaned.page_content
    assert "https://example.test  x" in cleaned.page_content
    assert f"knowledge-attachment://{ref}" in cleaned.page_content
    assert all(s.end <= len(cleaned.page_content) for s in cleaned.source_spans)


@pytest.mark.parametrize(
    "overrides",
    [
        {"size": 199},
        {"size": 4001},
        {"overlap": 501},
        {"mode": "parent_child", "child_size": 99},
        {"mode": "parent_child", "child_size": 2001},
        {"mode": "parent_child", "child_size": 1000},
    ],
)
def test_token_profile_limits_are_validated(overrides):
    with pytest.raises((ValueError, ExtractionError)):
        split([make_document("text")], **overrides)


def test_marked_word_table_rows_consume_header_once_and_keep_physical_cells(tmp_path):
    from actweave_knowledge.extraction.builtin.word_extractor import WordExtractor
    from docx import Document as WordDocument
    from docx.oxml import OxmlElement
    from parsing_test_helpers import make_context

    path = tmp_path / "table.docx"
    word = WordDocument()
    table = word.add_table(rows=1, cols=2)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    table.rows[0].cells[0].text = "编号"
    table.rows[0].cells[1].text = "详情"
    for i in range(3):
        table.add_row().cells[0].text = f"ID{i}"
        table.rows[-1].cells[1].text = "待处理" * 120
    word.save(path)
    documents = WordExtractor().parse_docx(path, make_context(tmp_path / "work"))
    assert documents[0].kind == "table_header"
    drafts = split(documents, size=200, overlap=50, child_size=100)
    assert_budgets(drafts)
    assert all("编号" in d.content and "详情" in d.content for d in drafts)
    assert sum(d.content.count("待处理") for d in drafts) == 360
    assert all(len({s.location.get("row") for s in d.source_spans if s.role == "source" and s.location.get("row", 1) > 1}) <= 1 for d in drafts)
    source_headers = [s for d in drafts for s in d.source_spans if s.block_id.startswith("table_header:") and s.role == "source"]
    assert {s.location["column"] for s in source_headers} == {1, 2}
    assert len(source_headers) == 2


def test_long_csv_fields_repeat_labels_but_not_values():
    from actweave_knowledge.extraction.contracts import HeaderRule
    from actweave_knowledge.extraction.tabular import rows_to_documents

    docs = rows_to_documents([(1, 1, ["编号", "详情"]), (2, 2, ["00123", "工单" * 400])], sheet=None, rule=HeaderRule(mode="explicit", row=1))
    drafts = split(docs, size=200, overlap=100, child_size=100)
    assert_budgets(drafts)
    value_drafts = [d for d in drafts if "工单" in d.content]
    assert len(value_drafts) > 1
    assert all("详情:" in d.content for d in value_drafts)
    assert "".join(d.content[s.start : s.end] for d in value_drafts for s in d.source_spans if s.role == "source" and s.location["column"] == 2) == "工单" * 400
    assert sum(d.content.count("00123") for d in drafts) == 1
    assert all(s.location["row"] == 2 for d in value_drafts for s in d.source_spans if s.role == "source")


def test_cleaning_external_link_keeps_visible_label_and_image_occurrence_offsets():
    from actweave_knowledge.ingestion.cleaner import clean_documents

    ref = "b" * 64
    text = f"说明 [操作手册](https://example.test/manual?q=1) 邮箱 a@b.test\n\n![原图](knowledge-attachment://{ref})"
    start = text.index("![")
    image = AttachmentOccurrence(ref=ref, alt_text="原图", source=SourceSpan(block_id="image", start=start, end=len(text), location={"paragraph": 2}))
    doc = make_document(text).model_copy(update={"attachments": (image,)})
    clean = clean_documents((doc,), remove_extra_spaces=True, remove_urls_emails=True)[0]
    assert "操作手册" in clean.page_content and "https://" not in clean.page_content
    assert "a@b.test" not in clean.page_content
    assert clean.page_content[clean.attachments[0].source.start : clean.attachments[0].source.end] == text[start:]


def test_ordinary_overlap_is_whole_units_without_carryover_only_segment():
    paragraphs = [f"段{i} " + "normal " * 45 for i in range(5)]
    drafts = split([make_document("\n\n".join(paragraphs))], size=200, overlap=60, child_size=100)
    assert_budgets(drafts)
    assert len(drafts) == 2
    assert [d.content.count("段3") for d in drafts] == [1, 1]
    assert sum(d.content.count("段4") for d in drafts) == 1


def test_character_profile_keeps_old_size_units_and_typed_children():
    doc = make_document("甲一。甲二。甲三。甲四。")
    drafts = split([doc], unit="character", mode="parent_child", size=7, overlap=0, child_size=4, separator="。", child_separator="。")
    assert [d.content for d in drafts] == ["甲一。甲二。", "甲三。甲四。"]
    assert [[c.content for c in d.children] for d in drafts] == [["甲一。", "甲二。"], ["甲三。", "甲四。"]]


def test_image_at_last_chunk_boundary_stays_bound_to_source_text():
    ref = "c" * 64
    text = "word " * 185 + f"\n\n![示意](knowledge-attachment://{ref})"
    start = text.index("![")
    occurrence = AttachmentOccurrence(ref=ref, alt_text="示意", source=SourceSpan(block_id="image", start=start, end=len(text), location={"paragraph": 1}))
    doc = make_document(text).model_copy(update={"attachments": (occurrence,)})
    drafts = split([doc], size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    images = [(d, item) for d in drafts for item in d.attachments]
    assert len(images) == 1
    draft, image = images[0]
    assert draft.content[image.source.start : image.source.end] == text[start:]
    assert "word" in draft.index_text
    assert sum(d.content.count("word") for d in drafts) == 185


def test_long_parent_children_keep_fences_and_table_fields():
    text = "```cpp\n" + "Map<K,V> value;\n" * 140 + "```"
    drafts = split([make_document(text)], mode="parent_child", size=400, overlap=0, child_size=100)
    assert_budgets(drafts, size=400)
    children = [c for d in drafts for c in d.children]
    assert all(c.content.startswith("```cpp\n") and c.content.endswith("```") for c in children)
    assert sum(c.content.count("Map<K,V> value;") for c in children) == 140


def test_trailing_image_does_not_break_previous_code_fences():
    from markdown_it import MarkdownIt

    ref = "d" * 64
    text = "```cpp\n" + "List<int> a;\n" * 35 + f"```\n\n![图](knowledge-attachment://{ref})"
    start = text.index("![")
    image = AttachmentOccurrence(ref=ref, alt_text="图", source=SourceSpan(block_id="image", start=start, end=len(text), location={"paragraph": 1}))
    drafts = split([make_document(text).model_copy(update={"attachments": (image,)})], size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    assert sum(d.content.count("List<int> a;") for d in drafts) == 35
    assert sum(len(d.attachments) for d in drafts) == 1
    for draft in drafts:
        tokens = MarkdownIt().parse(draft.content)
        assert sum(t.content.count("List<int> a;") for t in tokens if t.type == "fence") == draft.content.count("List<int> a;")
        assert draft.content.count("```") % 2 == 0


def test_parent_child_long_table_fields_keep_every_value_and_column_context():
    text = "| 编号 | 说明 |\n| --- | --- |\n| A123 | " + "detail " * 500 + " |"
    drafts = split([make_document(text, location={"table": 1, "row": 2})], mode="parent_child", size=400, overlap=0, child_size=100)
    assert_budgets(drafts, size=400)
    children = [c for d in drafts for c in d.children]
    assert sum(c.content.count("detail") for c in children) == 500
    assert all("说明" in c.content for c in children if "detail" in c.content)
    assert sum(c.content.count("A123") for c in children) == 1


def test_parent_hard_limit_stops_before_later_content(monkeypatch):
    from actweave_knowledge.contracts import KnowledgeError

    monkeypatch.setattr(splitter, "KNOWLEDGE_MAX_SEGMENTS_PER_DOCUMENT", 5)
    documents = tuple(make_document("record", location={"page": i}) for i in range(1, 7))
    assert len(split(documents[:5])) == 5
    too_late = make_document("# " + "huge " * 2000, location={"page": 7})
    with pytest.raises(KnowledgeError) as caught:
        split((*documents, too_late))
    assert caught.value.code == "KNOWLEDGE_QUOTA_EXCEEDED"


def test_vector_hard_limit_counts_children_separately_from_parents(monkeypatch):
    from actweave_knowledge.contracts import KnowledgeError

    monkeypatch.setattr(splitter, "KNOWLEDGE_MAX_SEGMENTS_PER_DOCUMENT", 4)
    documents = tuple(make_document("alpha " * 60 + "\n\n" + "beta " * 60, location={"page": i}) for i in range(1, 4))
    drafts = split(documents[:2], size=200, overlap=0, child_size=100, mode="parent_child")
    assert len(drafts) == 2 and sum(len(d.children) for d in drafts) == 4
    with pytest.raises(KnowledgeError) as caught:
        split(documents, size=200, overlap=0, child_size=100, mode="parent_child")
    assert caught.value.code == "KNOWLEDGE_QUOTA_EXCEEDED"


def test_review_heading_only_extracted_section_keeps_physical_heading():
    from actweave_knowledge.extraction.builtin.markdown_extractor import MarkdownExtractor

    documents = MarkdownExtractor().markdown_to_tups("# First\n\n# Second\n\nbody")
    assert len(documents) == 2
    drafts = split(documents)
    source_text = "".join(d.content[s.start : s.end] for d in drafts for s in d.source_spans if s.role == "source")
    assert "# First" in source_text and "# Second" in source_text and "body" in source_text
    assert sum(d.content.count("# First") for d in drafts) == 1
    assert any(s.location.get("line") == 1 for d in drafts for s in d.source_spans if s.role == "source")


def test_skipped_heading_levels_keep_all_real_ancestors():
    from actweave_knowledge.extraction.builtin.markdown_extractor import MarkdownExtractor

    text = "# Root\n\n### Middle\n\n#### Leaf\n\n" + "body " * 400
    documents = MarkdownExtractor().markdown_to_tups(text)
    drafts = split(documents, size=200, overlap=0, child_size=100)
    leaf_drafts = [d for d in drafts if "body" in d.content]
    assert len(leaf_drafts) > 1
    assert all(d.content.startswith("# Root\n\n### Middle\n\n#### Leaf") for d in leaf_drafts)
    for draft in leaf_drafts:
        middle = [s for s in draft.source_spans if "### Middle" in draft.content[s.start : s.end]]
        assert middle and all(s.role == "context_prefix" for s in middle)
        assert all(s.location.get("line") == 3 for s in middle)


def test_native_markdown_heading_keeps_raw_path_only_when_it_matches():
    from actweave_knowledge.extraction.builtin.markdown_extractor import MarkdownExtractor
    from actweave_knowledge.ingestion.structure import structure_groups

    (document,) = MarkdownExtractor().markdown_to_tups("# Native *Markdown* &amp;\n\nbody")
    prefix, pieces, _ = next(structure_groups((document,)))
    assert prefix.content == "# Native *Markdown* &amp;"
    assert pieces[0].content == "body"
    mismatched = document.model_copy(update={"heading_path": ("untrusted metadata",)})
    prefix, pieces, _ = next(structure_groups((mismatched,)))
    assert prefix.content == "# Native *Markdown* &amp;"
    assert pieces[0].content == "body"


def test_actual_word_leading_image_retains_exact_occurrence(tmp_path):
    from actweave_knowledge.extraction.builtin.word_extractor import WordExtractor
    from docx import Document as WordDocument
    from parsing_test_helpers import make_context
    from PIL import Image

    image_path = tmp_path / "pixel.png"
    Image.new("RGB", (1, 1), "blue").save(image_path)
    word = WordDocument()
    paragraph = word.add_paragraph()
    paragraph.add_run().add_picture(str(image_path))
    paragraph.add_run().add_break()
    paragraph.add_run("word " * 185)
    path = tmp_path / "leading-image.docx"
    word.save(path)
    documents = WordExtractor().parse_docx(path, make_context(tmp_path / "work"))
    original = documents[0].attachments[0]
    original_ref = documents[0].page_content[original.source.start : original.source.end]
    assert original_ref.startswith("![图片](knowledge-attachment:")
    drafts = split(documents, mode="parent_child", size=200, overlap=0, child_size=100)
    assert_budgets(drafts)
    occurrences = [(d, image) for d in drafts for image in d.attachments]
    assert len(occurrences) == 1
    draft, occurrence = occurrences[0]
    assert draft.content[occurrence.source.start : occurrence.source.end] == original_ref
    assert occurrence.ref == original.ref and occurrence.source.location == original.source.location
    assert draft.content.index(original_ref) < draft.content.index("word")
    assert sum(d.content.count("word") for d in drafts) == 185
    assert all("word" in d.index_text for d in drafts)
    assert sum(c.content.count("word") for d in drafts for c in d.children) == 185


def test_skipped_level_sibling_does_not_inherit_previous_branch():
    from actweave_knowledge.extraction.builtin.markdown_extractor import MarkdownExtractor

    text = "# Root\n\n### Middle\n\n#### Leaf\n\nleaf body\n\n### Sibling\n\nsibling body"
    documents = MarkdownExtractor().markdown_to_tups(text)
    drafts = split(documents)
    sibling = next(d for d in drafts if "sibling body" in d.content)
    assert sibling.content.startswith("# Root\n\n### Sibling")
    assert "Middle" not in sibling.content and "Leaf" not in sibling.content
    source = "".join(d.content[s.start : s.end] for d in drafts for s in d.source_spans if s.role == "source")
    assert source.count("# Root") == 1 and source.count("### Middle") == 1 and source.count("#### Leaf") == 1


def test_review_leading_image_reserves_budget_without_breaking_following_code():
    from markdown_it import MarkdownIt

    ref = "e" * 64
    image_text = f"![图片](knowledge-attachment:{ref})"
    text = image_text + "\n\n```cpp\n" + "List<int> a;\n" * 36 + "```"
    occurrence = AttachmentOccurrence(ref=ref, alt_text="图片", source=SourceSpan(block_id="image", start=0, end=len(image_text), location={"paragraph": 1}))
    drafts = split([make_document(text).model_copy(update={"attachments": (occurrence,)})], size=200, overlap=0, child_size=100, mode="parent_child")
    assert_budgets(drafts)
    assert sum(len(d.attachments) for d in drafts) == 1
    assert drafts[0].content.startswith(image_text)
    assert sum(d.content.count("List<int> a;") for d in drafts) == 36
    for value in [*drafts, *(c for d in drafts for c in d.children)]:
        assert value.content.count("```") % 2 == 0
        assert sum(t.content.count("List<int> a;") for t in MarkdownIt().parse(value.content) if t.type == "fence") == value.content.count("List<int> a;")


def test_review_impossible_image_and_atomic_text_budget_fails_without_partial_result():
    ref = "f" * 64
    image_text = f"![图片](knowledge-attachment:{ref})"
    text = image_text + "\n" + "[actual source](https://example.test/" + "0123456789abcdef" * 100 + ")"
    occurrence = AttachmentOccurrence(ref=ref, alt_text="图片", source=SourceSpan(block_id="image", start=0, end=len(image_text), location={"paragraph": 1}))
    with pytest.raises(ExtractionError) as caught:
        split([make_document(text).model_copy(update={"attachments": (occurrence,)})], size=200, overlap=0, child_size=100)
    assert caught.value.reason_code == "ATOMIC_CONTENT_EXCEEDS_BUDGET"


def test_real_markdown_placeholder_preserves_indented_code_payload_and_sources():
    from actweave_knowledge.extraction.builtin.markdown_extractor import MarkdownExtractor
    from markdown_it import MarkdownIt

    source = "![image](https://example.test/image.png)\n\n" + "    total += 1\n" * 33
    documents = MarkdownExtractor().markdown_to_tups(source)
    canonical = documents[0].page_content
    assert canonical.startswith("（外部图片未获取：image）")
    expected = "".join(token.content for token in MarkdownIt().parse(canonical) if token.type == "code_block")
    assert expected == "total += 1\n" * 33
    drafts = split(documents, size=200, overlap=0, child_size=100, mode="parent_child")
    assert_budgets(drafts)
    parent_payload = "".join(token.content for draft in drafts for token in MarkdownIt().parse(draft.content) if token.type in {"code_block", "fence"})
    assert parent_payload == expected
    compile(parent_payload, "chunked_source", "exec")
    groups = [drafts, [child for draft in drafts for child in draft.children]]
    for group in groups:
        payload = "".join(token.content for value in group for token in MarkdownIt().parse(value.content) if token.type in {"code_block", "fence"})
        assert payload == expected
        covered = {}
        for value in group:
            for span in value.source_spans:
                if span.role == "source" and span.location.get("line", 0) >= 3:
                    covered.setdefault(span.location["line"], []).append(value.content[span.start : span.end])
        assert set(covered) == set(range(3, 36))
        assert all("".join(parts).rstrip("\n") == "total += 1" for parts in covered.values())


@pytest.mark.parametrize("indent", ["    ", "\t"])
def test_code_conversion_removes_only_markdown_indent(indent):
    from actweave_knowledge.extraction.builtin.markdown_extractor import MarkdownExtractor
    from markdown_it import MarkdownIt

    source = (indent + "if ready:\n" + indent + "    total += 1\n") * 35
    documents = MarkdownExtractor().markdown_to_tups(source)
    expected = "".join(t.content for document in documents for t in MarkdownIt().parse(document.page_content) if t.type == "code_block")
    assert expected == "if ready:\n    total += 1\n" * 35
    drafts = split(documents, size=200, overlap=0, child_size=100, mode="parent_child")
    assert_budgets(drafts)
    values = [child for draft in drafts for child in draft.children]
    actual = "".join(t.content for value in values for t in MarkdownIt().parse(value.content) if t.type in {"code_block", "fence"})
    assert actual == expected
    compile(actual, "chunked_nested_source", "exec")
    assert {s.location["line"] for value in values for s in value.source_spans if s.role == "source"} == set(range(1, 71))

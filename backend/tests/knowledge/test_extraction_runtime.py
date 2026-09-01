from __future__ import annotations

import os
import sys
from pathlib import Path

import pytest
from actweave_knowledge.extraction import runtime
from actweave_knowledge.extraction.contracts import ExtractionError, ExtractionLimits
from parsing_test_helpers import make_setting


async def _guard():
    return None


async def _no_asset(asset):
    raise AssertionError("fixture cannot produce an asset")


@pytest.mark.asyncio
async def test_timeout_reaps_child_before_return(tmp_path, monkeypatch):
    source = tmp_path / "source.txt"
    source.write_text("hello", encoding="utf-8")
    pidfile = tmp_path / "child.pid"
    fixture = Path(__file__).parent / "fixtures" / "parser_child_hang.py"
    monkeypatch.setattr(runtime, "sandbox_command", lambda command, *, work_dir: [sys.executable, str(fixture), str(pidfile)])
    with pytest.raises(ExtractionError) as caught:
        await runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=1, guard=_guard, on_asset=_no_asset)
    assert caught.value.reason_code == "PARSER_TIMEOUT"
    with pytest.raises(ProcessLookupError):
        os.kill(int(pidfile.read_text()), 0)
    assert (tmp_path / "work").exists()


@pytest.mark.asyncio
async def test_parser_slots_reject_busy_without_queueing():
    slots = runtime.ParserSlots(1)
    async with slots:
        with pytest.raises(ExtractionError, match="解析繁忙") as caught:
            async with slots:
                raise AssertionError("second parser must not start")
        assert slots.active == 1
    assert caught.value.reason_code == "PARSER_BUSY"
    assert slots.active == 0


@pytest.mark.asyncio
async def test_real_sandbox_text_roundtrip(tmp_path):
    source = tmp_path / "source.txt"
    source.write_text("hello local parser", encoding="utf-8")
    result = await runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=30, guard=_guard, on_asset=_no_asset)
    assert "hello local parser" in "\n".join(doc.page_content for doc in result.documents)
    assert result.documents[0].source_spans


def test_environment_does_not_inherit_secrets(tmp_path, monkeypatch):
    monkeypatch.setenv("DATABASE_URL", "secret-db")
    monkeypatch.setenv("HTTP_PROXY", "secret-proxy")
    monkeypatch.setenv("SOME_FUTURE_SECRET", "secret")
    environment = runtime.parser_environment(tmp_path)
    assert not {"DATABASE_URL", "HTTP_PROXY", "SOME_FUTURE_SECRET"} & environment.keys()
    assert environment["HOME"] == str(tmp_path.resolve() / "child")


def _attack(tmp_path, monkeypatch, mode):
    source = tmp_path / "source.txt"
    source.write_text("hello")
    pidfile = tmp_path / "child.pid"
    fixture = Path(__file__).parent / "fixtures" / "parser_child_hang.py"
    monkeypatch.setattr(runtime, "sandbox_command", lambda command, *, work_dir: [sys.executable, str(fixture), str(pidfile), mode, command[-1]])
    return make_setting(source), pidfile


@pytest.mark.asyncio
@pytest.mark.parametrize("mode", ["bad:unknown", "bad:huge", "bad:truncated", "bad:duplicate", "bad:empty", "bad:secret", "asset_eof", "asset_missing", "result:source", "result:profile", "result:twice"])
async def test_bounded_protocol_rejects_without_success(tmp_path, monkeypatch, mode):
    setting, pidfile = _attack(tmp_path, monkeypatch, mode)

    async def asset(value):
        assert value.relative_path.startswith("received/")

    with pytest.raises(ExtractionError) as error:
        await runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=10, guard=_guard, on_asset=asset)
    assert error.value.reason_code == "PARSER_OUTPUT_INVALID"
    with pytest.raises(ProcessLookupError):
        os.kill(int(pidfile.read_text()), 0)


@pytest.mark.asyncio
async def test_cancel_reaps_before_settling_started_callback(tmp_path, monkeypatch):
    import asyncio

    setting, pidfile = _attack(tmp_path, monkeypatch, "asset")
    started, released, settled, reaped = (asyncio.Event() for _ in range(4))
    stop = runtime.stop_process_group

    async def stopped(process):
        await stop(process)
        reaped.set()

    monkeypatch.setattr(runtime, "stop_process_group", stopped)

    async def asset(value):
        started.set()
        await released.wait()
        assert reaped.is_set()
        assert (tmp_path / "work" / value.relative_path).exists()
        settled.set()

    task = asyncio.create_task(runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=10, guard=_guard, on_asset=asset))
    await asyncio.wait_for(started.wait(), 5)
    task.cancel()
    await asyncio.wait_for(reaped.wait(), 5)
    task.cancel()  # a second cancellation during cleanup must not abandon I/O
    assert not task.done()
    with pytest.raises(ProcessLookupError):
        os.kill(int(pidfile.read_text()), 0)
    released.set()
    with pytest.raises(asyncio.CancelledError):
        await asyncio.wait_for(task, 5)
    assert settled.is_set()
    assert not (tmp_path / "work/child/ack").exists()


@pytest.mark.asyncio
async def test_cancel_joins_inflight_receiver_thread(tmp_path, monkeypatch):
    import asyncio
    import threading

    setting, _ = _attack(tmp_path, monkeypatch, "asset")
    started, released = threading.Event(), threading.Event()
    reaped = asyncio.Event()
    stop, receive = runtime.stop_process_group, runtime.receive_asset

    async def stopped(process):
        await stop(process)
        reaped.set()

    def receiver(*args, **kwargs):
        started.set()
        assert released.wait(10)
        return receive(*args, **kwargs)

    monkeypatch.setattr(runtime, "stop_process_group", stopped)
    monkeypatch.setattr(runtime, "receive_asset", receiver)
    task = asyncio.create_task(runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=10, guard=_guard, on_asset=_no_asset))
    assert await asyncio.to_thread(started.wait, 5)
    task.cancel()
    await asyncio.wait_for(reaped.wait(), 5)
    assert not task.done()
    released.set()
    with pytest.raises(asyncio.CancelledError):
        await asyncio.wait_for(task, 5)
    assert list((tmp_path / "work/received").glob("*.png"))


@pytest.mark.asyncio
async def test_guard_revoked_after_callback_blocks_ack_and_preserves_error(tmp_path, monkeypatch):
    setting, pidfile = _attack(tmp_path, monkeypatch, "asset")
    revoked = False
    error = RuntimeError("test authority revoked")

    async def guard():
        if revoked:
            raise error

    async def asset(value):
        nonlocal revoked
        revoked = True

    with pytest.raises(RuntimeError) as caught:
        await runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=10, guard=guard, on_asset=asset)
    assert caught.value is error
    assert not (tmp_path / "work/child/ack").exists()
    with pytest.raises(ProcessLookupError):
        os.kill(int(pidfile.read_text()), 0)


@pytest.mark.asyncio
async def test_disk_growth_kills_real_child(tmp_path, monkeypatch):
    setting, pidfile = _attack(tmp_path, monkeypatch, "budget")
    with pytest.raises(ExtractionError) as caught:
        await runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(max_work_dir_bytes=65536), timeout_seconds=10, guard=_guard, on_asset=_no_asset)
    assert caught.value.reason_code == "PARSER_WORK_DIR_LIMIT_EXCEEDED"
    with pytest.raises(ProcessLookupError):
        os.kill(int(pidfile.read_text()), 0)


@pytest.mark.asyncio
async def test_callback_timeout_error_is_not_changed_to_parser_timeout(tmp_path, monkeypatch):
    setting, _ = _attack(tmp_path, monkeypatch, "asset")
    error = TimeoutError("host object I/O timeout")

    async def asset(value):
        raise error

    with pytest.raises(TimeoutError) as caught:
        await runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=10, guard=_guard, on_asset=asset)
    assert caught.value is error


def test_os_sandbox_denies_real_live_connection_and_host_files(tmp_path):
    import json
    import socket
    import subprocess

    from actweave_knowledge.extraction.sandbox import parser_environment, sandbox_command

    root = tmp_path.resolve()
    for name in ("source", "child", "received", "outside"):
        (root / name).mkdir()
    (root / "source/input.txt").write_text("safe fixture")
    (root / "outside/sentinel").write_text("private fixture")
    (root / "received/sentinel").write_text("parent-owned image bytes")
    code = """import json,pathlib,socket,sys
root=pathlib.Path(sys.argv[1]); result={}
for name,path,write in [
 ("child",root/"child/allowed",True),("received",root/"received/denied",True),
 ("outside_write",root/"outside/denied",True),("source",root/"source/input.txt",False),
 ("outside_read",root/"outside/sentinel",False),("received_read",root/"received/sentinel",False),
 ("source_write",root/"source/input.txt",True),("root_write",root/"denied",True)]:
 try:
  path.write_text("ok") if write else path.read_text(); result[name]="allowed"
 except OSError: result[name]="denied"
s=socket.socket(); s.settimeout(1)
try: s.connect(("127.0.0.1",int(sys.argv[2]))); result["connect"]="allowed"
except OSError as error: result["connect"]="denied"; result["connect_errno"]=error.errno
print(json.dumps(result))
"""
    with socket.socket() as listener:
        listener.bind(("127.0.0.1", 0))
        listener.listen(4)
        with socket.create_connection(listener.getsockname()):
            control, _ = listener.accept()
            control.close()  # prove listener is live before sandbox launch
        process = subprocess.run(
            sandbox_command([sys.executable, "-s", "-B", "-c", code, str(root), str(listener.getsockname()[1])], work_dir=root), cwd=root / "child", env=parser_environment(root), capture_output=True, text=True, timeout=15
        )
        assert process.returncode == 0, process.stderr
        result = json.loads(process.stdout)
        assert result["child"] == result["source"] == "allowed"
        assert all(result[name] == "denied" for name in ("received", "outside_write", "outside_read", "received_read", "source_write", "root_write", "connect"))
        if sys.platform == "darwin":
            assert result["connect_errno"] == 1  # actual Seatbelt EPERM
        listener.setblocking(False)
        with pytest.raises(BlockingIOError):
            listener.accept()  # accept count zero, not a DNS failure simulation


@pytest.mark.asyncio
async def test_in_work_source_is_copied_without_mutating_caller_path(tmp_path):
    work = tmp_path / "work"
    work.mkdir()
    source = work / "caller.txt"
    source.write_text("caller source remains readable")
    await runtime.run_extraction(make_setting(source), work_dir=work, limits=ExtractionLimits(), timeout_seconds=30, guard=_guard, on_asset=_no_asset)
    assert source.read_text() == "caller source remains readable"
    assert (work / "source/input.txt").read_bytes() == source.read_bytes()


@pytest.mark.asyncio
async def test_sandbox_launch_failure_is_fail_closed(tmp_path, monkeypatch):
    source = tmp_path / "source.txt"
    source.write_text("hello")
    monkeypatch.setattr(runtime, "sandbox_command", lambda command, *, work_dir: [sys.executable, "-c", "raise SystemExit(42)"])
    with pytest.raises(ExtractionError) as caught:
        await runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=10, guard=_guard, on_asset=_no_asset)
    assert caught.value.reason_code == "PARSER_SANDBOX_UNAVAILABLE"


@pytest.mark.asyncio
async def test_exited_leader_does_not_leave_running_descendant(tmp_path, monkeypatch):
    import subprocess

    setting, pidfile = _attack(tmp_path, monkeypatch, "descendant")
    with pytest.raises(ExtractionError) as caught:
        await runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=1, guard=_guard, on_asset=_no_asset)
    assert caught.value.reason_code == "PARSER_TIMEOUT"
    descendant = int(Path(str(pidfile) + ".descendant").read_text())
    if sys.platform.startswith("linux"):
        try:
            stat = Path(f"/proc/{descendant}/stat").read_text(encoding="utf-8")
        except FileNotFoundError:
            return
        _, separator, fields = stat.rpartition(")")
        assert separator and fields.split()[0] == "Z"
    else:
        status = subprocess.run(["ps", "-o", "stat=", "-p", str(descendant)], capture_output=True, text=True)
        assert not status.stdout.strip() or status.stdout.strip().startswith("Z")


def test_linux_group_scan_fails_closed_when_initial_stat_is_unreadable(tmp_path, monkeypatch):
    proc = tmp_path / "proc"
    stat_path = proc / "123/stat"
    stat_path.parent.mkdir(parents=True)
    stat_path.write_text("unused", encoding="utf-8")
    scandir = os.scandir
    read_text = runtime.Path.read_text

    monkeypatch.setattr(runtime.os, "scandir", lambda path: scandir(proc))

    def denied(path, *args, **kwargs):
        if path == stat_path:
            raise PermissionError("denied synthetic proc stat")
        return read_text(path, *args, **kwargs)

    monkeypatch.setattr(runtime.Path, "read_text", denied)
    with pytest.raises(RuntimeError, match="process-group state is unavailable"):
        runtime._linux_live_process_group_pidfds(123)


def test_linux_group_scan_fails_closed_when_revalidation_is_unreadable(tmp_path, monkeypatch):
    proc = tmp_path / "proc"
    stat_path = proc / "123/stat"
    stat_path.parent.mkdir(parents=True)
    stat_path.write_text("unused", encoding="utf-8")
    scandir = os.scandir
    read_text = runtime.Path.read_text
    reads = 0
    read_handle, write_handle = os.pipe()

    monkeypatch.setattr(runtime.os, "scandir", lambda path: scandir(proc))
    monkeypatch.setattr(runtime.os, "pidfd_open", lambda pid: read_handle, raising=False)

    def changing_stat(path, *args, **kwargs):
        nonlocal reads
        if path != stat_path:
            return read_text(path, *args, **kwargs)
        reads += 1
        if reads == 1:
            return "123 (fixture) S 1 123 123"
        raise PermissionError("denied synthetic proc stat revalidation")

    monkeypatch.setattr(runtime.Path, "read_text", changing_stat)
    try:
        with pytest.raises(RuntimeError, match="process-group state is unavailable"):
            runtime._linux_live_process_group_pidfds(123)
        with pytest.raises(OSError):
            os.fstat(read_handle)
    finally:
        os.close(write_handle)


@pytest.mark.asyncio
@pytest.mark.parametrize("registered_before_failure", [0, 1])
async def test_linux_group_wait_releases_handles_when_reader_registration_fails(monkeypatch, registered_before_failure):
    import asyncio

    loop = asyncio.get_running_loop()
    add_reader = loop.add_reader
    remove_reader = loop.remove_reader
    pipes = [os.pipe() for _ in range(registered_before_failure + 1)]
    handles = tuple(read_handle for read_handle, _ in pipes)
    registrations = 0
    error = RuntimeError("synthetic reader registration failure")

    monkeypatch.setattr(runtime, "_linux_live_process_group_pidfds", lambda process_group_id: handles)

    def failing_add_reader(handle, callback):
        nonlocal registrations
        if registrations == registered_before_failure:
            raise error
        registrations += 1
        add_reader(handle, callback)

    monkeypatch.setattr(loop, "add_reader", failing_add_reader)
    try:
        with pytest.raises(RuntimeError) as caught:
            await runtime._wait_linux_process_group_quiescent(123)
        assert caught.value is error
        for handle in handles:
            with pytest.raises(OSError):
                os.fstat(handle)
        assert all(not remove_reader(handle) for handle in handles[:registered_before_failure])
    finally:
        for read_handle, write_handle in pipes:
            remove_reader(read_handle)
            try:
                os.close(read_handle)
            except OSError:
                pass
            os.close(write_handle)


@pytest.mark.asyncio
async def test_image_only_document_returns_without_ocr(tmp_path):
    from docx import Document
    from PIL import Image

    png = tmp_path / "red.png"
    Image.new("RGB", (3, 2), "red").save(png)
    document = Document()
    document.add_picture(str(png))
    source = tmp_path / "image-only.docx"
    document.save(source)
    assets = []

    async def asset(value):
        assets.append(value)

    result = await runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=30, guard=_guard, on_asset=asset)
    assert len(assets) == len(result.attachments) == 1
    assert all(document.attachments for document in result.documents)
    assert all(document.page_content.startswith("![") for document in result.documents)


@pytest.mark.asyncio
async def test_timeout_settles_callback_after_process_reap(tmp_path, monkeypatch):
    import asyncio

    setting, _ = _attack(tmp_path, monkeypatch, "asset")
    started, released, reaped = (asyncio.Event() for _ in range(3))
    stop = runtime.stop_process_group

    async def stopped(process):
        await stop(process)
        reaped.set()

    monkeypatch.setattr(runtime, "stop_process_group", stopped)

    async def asset(value):
        started.set()
        await released.wait()
        assert reaped.is_set()

    task = asyncio.create_task(runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=1, guard=_guard, on_asset=asset))
    await asyncio.wait_for(started.wait(), 5)
    await asyncio.wait_for(reaped.wait(), 5)
    assert not task.done()
    released.set()
    with pytest.raises(ExtractionError) as caught:
        await asyncio.wait_for(task, 5)
    assert caught.value.reason_code == "PARSER_TIMEOUT"


@pytest.mark.asyncio
async def test_process_group_error_still_settles_started_callback_before_return(tmp_path, monkeypatch):
    import asyncio

    setting, _ = _attack(tmp_path, monkeypatch, "asset")
    started, released, stop_failed, settled = (asyncio.Event() for _ in range(4))
    stop = runtime.stop_process_group
    error = RuntimeError("synthetic process-group state unavailable")

    async def failed_stop(process):
        await stop(process)
        stop_failed.set()
        raise error

    monkeypatch.setattr(runtime, "stop_process_group", failed_stop)

    async def asset(value):
        started.set()
        await released.wait()
        settled.set()

    task = asyncio.create_task(runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=1, guard=_guard, on_asset=asset))
    await asyncio.wait_for(started.wait(), 5)
    await asyncio.wait_for(stop_failed.wait(), 5)
    try:
        with pytest.raises(TimeoutError):
            await asyncio.wait_for(asyncio.shield(task), 0.1)
    finally:
        released.set()
    with pytest.raises(RuntimeError) as caught:
        await asyncio.wait_for(task, 5)
    assert caught.value is error
    assert settled.is_set()


def test_linux_launcher_has_no_bare_process_fallback(tmp_path, monkeypatch):
    from actweave_knowledge.extraction import sandbox

    monkeypatch.setattr(sandbox.sys, "platform", "linux")
    monkeypatch.setattr(sandbox.shutil, "which", lambda *args, **kwargs: None)
    with pytest.raises(ExtractionError) as caught:
        sandbox.sandbox_command([sys.executable, "-m", "actweave_knowledge.extraction.child"], work_dir=tmp_path)
    assert caught.value.reason_code == "PARSER_SANDBOX_UNAVAILABLE"
    monkeypatch.setattr(sandbox.shutil, "which", lambda *args, **kwargs: "/usr/bin/bwrap")
    command = sandbox.sandbox_command([sys.executable, "-m", "actweave_knowledge.extraction.child", "--output-fd", "8"], work_dir=tmp_path)
    assert command[0] == "/usr/bin/bwrap"
    assert {"--unshare-user", "--unshare-net", "--unshare-pid", "--die-with-parent", "--proc", "--dev", "--remount-ro"} <= set(command)
    writable = [command[index + 1] for index, word in enumerate(command) if word == "--bind"]
    assert writable == [str(tmp_path.resolve() / "child")]
    assert not any(command[index : index + 3] == ["--ro-bind", "/", "/"] for index in range(len(command)))


@pytest.mark.asyncio
async def test_cancellation_arriving_during_success_cleanup_is_not_swallowed(tmp_path, monkeypatch):
    import asyncio

    source = tmp_path / "source.txt"
    source.write_text("hello")
    stopping, release = asyncio.Event(), asyncio.Event()
    stop = runtime.stop_process_group

    async def stopped(process):
        stopping.set()
        await release.wait()
        await stop(process)

    monkeypatch.setattr(runtime, "stop_process_group", stopped)
    task = asyncio.create_task(runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=30, guard=_guard, on_asset=_no_asset))
    await asyncio.wait_for(stopping.wait(), 5)
    task.cancel()
    release.set()
    with pytest.raises(asyncio.CancelledError):
        await asyncio.wait_for(task, 5)


def test_space_monitor_tolerates_conversion_file_disappearing(tmp_path, monkeypatch):
    transient = tmp_path / "conversion.tmp"
    transient.write_bytes(b"temporary")
    original = os.stat

    def remove_then_stat(path, *args, **kwargs):
        if Path(path) == transient:
            transient.unlink(missing_ok=True)
        return original(path, *args, **kwargs)

    monkeypatch.setattr(os, "stat", remove_then_stat)
    runtime._check_work_dir(tmp_path, ExtractionLimits())


@pytest.mark.asyncio
async def test_revocation_during_source_preparation_prevents_spawn(tmp_path, monkeypatch):
    import asyncio
    import threading

    source = tmp_path / "source.txt"
    source.write_text("hello")
    prepared, release = threading.Event(), threading.Event()
    original_prepare = runtime._prepare
    original_spawn = asyncio.create_subprocess_exec
    calls = []
    revoked = False
    error = RuntimeError("authority revoked during source preparation")

    def prepare(*args):
        result = original_prepare(*args)
        prepared.set()
        assert release.wait(10)
        return result

    async def spawn(*args, **kwargs):
        calls.append(True)
        return await original_spawn(*args, **kwargs)

    async def guard():
        if revoked:
            raise error

    monkeypatch.setattr(runtime, "_prepare", prepare)
    monkeypatch.setattr(asyncio, "create_subprocess_exec", spawn)
    task = asyncio.create_task(runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=30, guard=guard, on_asset=_no_asset))
    assert await asyncio.to_thread(prepared.wait, 5)
    revoked = True
    release.set()
    with pytest.raises(RuntimeError) as caught:
        await asyncio.wait_for(task, 5)
    assert caught.value is error
    assert calls == []


@pytest.mark.parametrize("resource_kind", ["pandoc", "nlp"])
@pytest.mark.parametrize("damage", ["missing", "tampered"])
def test_real_sandbox_missing_or_tampered_resource_copy_fails_closed(tmp_path, resource_kind, damage):
    import json
    import shutil
    import subprocess

    from actweave_knowledge.extraction import runtime_resources
    from actweave_knowledge.extraction.sandbox import parser_environment, sandbox_command

    root = tmp_path.resolve()
    (root / "source").mkdir()
    (root / "child").mkdir()
    manifest = runtime_resources.runtime_manifest()
    logical = (
        "pypandoc-binary/pypandoc/files/pandoc" if resource_kind == "pandoc" else next(row["logical_name"] for row in manifest["resources"] if row["logical_name"].startswith("en-core-web-sm/") and row["logical_name"].endswith("meta.json"))
    )
    copy = root / "source/resource-copy"
    shutil.copyfile(runtime_resources._resource_path(logical), copy)
    if damage == "missing":
        copy.unlink()
    else:
        with copy.open("ab") as output:
            output.write(b"tampered")
    code = """import json,sys
from pathlib import Path
from actweave_knowledge.extraction import runtime_resources as r
from actweave_knowledge.extraction.contracts import ExtractionError
before=r.runtime_digest(); original=r._resource_path
r._resource_path=lambda name: Path(sys.argv[2]) if name==sys.argv[1] else original(name)
parser="unstructured.epub" if sys.argv[3]=="pandoc" else "unstructured.pptx"
result={"probe":r.probe_parser_resources(parser),"digest_changed":before!=r.runtime_digest()}
try: r.prepare_local_parser(parser); result["prepared"]=True
except ExtractionError as error: result["prepare_error"]=error.reason_code
result["partition_imported"]=any(name.startswith("unstructured.partition") for name in sys.modules)
print(json.dumps(result))
"""
    process = subprocess.run(sandbox_command([sys.executable, "-s", "-B", "-c", code, logical, str(copy), resource_kind], work_dir=root), cwd=root / "child", env=parser_environment(root), capture_output=True, text=True, timeout=20)
    assert process.returncode == 0, process.stderr
    result = json.loads(process.stdout)
    assert result == {"probe": "PARSER_DEPENDENCY_UNAVAILABLE", "digest_changed": True, "prepare_error": "PARSER_DEPENDENCY_UNAVAILABLE", "partition_imported": False}
    assert not copy.exists() if damage == "missing" else copy.exists()


@pytest.mark.asyncio
async def test_cancel_stops_new_callback_before_waiting_for_process_reap(tmp_path, monkeypatch):
    import asyncio
    import threading

    setting, _ = _attack(tmp_path, monkeypatch, "asset")
    receiver_started, release_receiver, receiver_done = (threading.Event() for _ in range(3))
    stop_started, allow_stop, asset_started = (asyncio.Event() for _ in range(3))
    receive, stop = runtime.receive_asset, runtime.stop_process_group

    def receiver(*args, **kwargs):
        receiver_started.set()
        assert release_receiver.wait(10)
        result = receive(*args, **kwargs)
        receiver_done.set()
        return result

    async def stopped(process):
        stop_started.set()
        await allow_stop.wait()
        await stop(process)

    async def asset(value):
        asset_started.set()

    monkeypatch.setattr(runtime, "receive_asset", receiver)
    monkeypatch.setattr(runtime, "stop_process_group", stopped)
    task = asyncio.create_task(runtime.run_extraction(setting, work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=30, guard=_guard, on_asset=asset))
    assert await asyncio.to_thread(receiver_started.wait, 5)
    task.cancel()
    await asyncio.wait_for(stop_started.wait(), 5)
    release_receiver.set()
    assert await asyncio.to_thread(receiver_done.wait, 5)
    # The receiver finishes while process termination is held at an explicit
    # barrier. No callback may begin in this interval; timeout observes absence.
    try:
        await asyncio.wait_for(asset_started.wait(), 0.1)
        dispatched = True
    except TimeoutError:
        dispatched = False
    finally:
        allow_stop.set()
        with pytest.raises(asyncio.CancelledError):
            await asyncio.wait_for(task, 5)
    assert not dispatched


@pytest.mark.asyncio
@pytest.mark.parametrize("error_type", [BrokenPipeError, ConnectionResetError])
@pytest.mark.parametrize("origin", ["on_asset", "post_asset_guard"])
async def test_real_sandbox_preserves_host_transport_exception_identity(tmp_path, monkeypatch, error_type, origin):
    from docx import Document
    from PIL import Image

    png = tmp_path / "red.png"
    Image.new("RGB", (3, 2), "red").save(png)
    document = Document()
    document.add_picture(str(png))
    source = tmp_path / "host-error.docx"
    document.save(source)
    error = error_type("host transport failed")
    received = []
    reaped = []
    stop = runtime.stop_process_group

    async def stopped(process):
        await stop(process)
        reaped.append(process.pid)

    async def guard():
        if origin == "post_asset_guard" and received:
            raise error

    async def on_asset(asset):
        assert asset.relative_path.startswith("received/")
        assert (tmp_path / "work" / asset.relative_path).is_file()
        received.append(asset)
        if origin == "on_asset":
            raise error

    monkeypatch.setattr(runtime, "stop_process_group", stopped)
    with pytest.raises(error_type) as caught:
        await runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=30, guard=guard, on_asset=on_asset)
    assert caught.value is error
    assert len(received) == len(reaped) == 1
    with pytest.raises(ProcessLookupError):
        os.kill(reaped[0], 0)


@pytest.mark.asyncio
async def test_real_child_protocol_ack_pipe_failure_is_parser_output_invalid(tmp_path, monkeypatch):
    import asyncio

    from docx import Document
    from PIL import Image

    png = tmp_path / "red.png"
    Image.new("RGB", (3, 2), "red").save(png)
    document = Document()
    document.add_picture(str(png))
    source = tmp_path / "closed-ack-pipe.docx"
    document.save(source)
    children = []
    spawn = asyncio.create_subprocess_exec

    async def spawned(*args, **kwargs):
        process = await spawn(*args, **kwargs)
        children.append(process)
        return process

    async def on_asset(asset):
        # Break the actual child protocol transport after receiving its valid
        # asset. This callback returns normally; only the subsequent ACK fails.
        assert len(children) == 1
        children[0].stdin.close()
        await children[0].stdin.wait_closed()

    monkeypatch.setattr(asyncio, "create_subprocess_exec", spawned)
    with pytest.raises(ExtractionError) as caught:
        await runtime.run_extraction(make_setting(source), work_dir=tmp_path / "work", limits=ExtractionLimits(), timeout_seconds=30, guard=_guard, on_asset=on_asset)
    assert caught.value.reason_code == "PARSER_OUTPUT_INVALID"
    assert children[0].returncode is not None
    with pytest.raises(ProcessLookupError):
        os.kill(children[0].pid, 0)
